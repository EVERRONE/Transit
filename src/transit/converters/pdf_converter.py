"""PDF to DOCX converter using pdf2docx."""

import logging
import os
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple

logger = logging.getLogger(__name__)


class PDFConverter:
    """
    Convert PDF files to DOCX format using pdf2docx.

    Handles conversion with quality validation and error reporting.
    """

    def __init__(
        self,
        start_page: int = 0,
        end_page: Optional[int] = None,
        multi_processing: bool = False,
        cpu_count: int = 0
    ):
        """
        Initialize PDF converter.

        Args:
            start_page: First page to convert (0-indexed)
            end_page: Last page to convert (None = all pages)
            multi_processing: Enable multi-processing for large PDFs
            cpu_count: Number of CPUs to use (0 = auto)
        """
        self.start_page = start_page
        self.end_page = end_page
        self.multi_processing = multi_processing
        self.cpu_count = cpu_count

        # Check if pdf2docx is available
        try:
            from pdf2docx import Converter
            self._converter_available = True
        except ImportError:
            self._converter_available = False
            logger.warning(
                "pdf2docx not installed. Install with: pip install pdf2docx"
            )

        logger.info(
            f"Initialized PDF converter: "
            f"pages={start_page}-{end_page or 'end'}, "
            f"multi_processing={multi_processing}"
        )

    def is_available(self) -> bool:
        """
        Check if PDF conversion is available.

        Returns:
            True if pdf2docx is installed
        """
        return self._converter_available

    def convert_pdf_to_docx(
        self,
        pdf_path: str,
        docx_path: Optional[str] = None,
        pages: Optional[List[int]] = None
    ) -> Tuple[bool, str, Dict[str, Any]]:
        """
        Convert PDF to DOCX.

        Args:
            pdf_path: Path to input PDF file
            docx_path: Path to output DOCX file (default: same name with .docx)
            pages: Optional list of specific pages to convert (0-indexed)

        Returns:
            Tuple of (success, output_path, conversion_stats)
        """
        if not self._converter_available:
            return (
                False,
                "",
                {"error": "pdf2docx not installed"}
            )

        # Validate input file
        pdf_path = Path(pdf_path)
        if not pdf_path.exists():
            return (
                False,
                "",
                {"error": f"PDF file not found: {pdf_path}"}
            )

        if not pdf_path.suffix.lower() == '.pdf':
            return (
                False,
                "",
                {"error": f"Not a PDF file: {pdf_path}"}
            )

        # Determine output path
        if not docx_path:
            docx_path = pdf_path.with_suffix('.docx')
        else:
            docx_path = Path(docx_path)

        logger.info(f"Converting PDF: {pdf_path} -> {docx_path}")

        try:
            from pdf2docx import Converter

            # Create converter
            cv = Converter(str(pdf_path))

            # Determine page range
            if pages:
                # Convert specific pages
                logger.info(f"Converting specific pages: {pages}")
                start = min(pages)
                end = max(pages) + 1
            else:
                # Use configured range
                start = self.start_page
                end = self.end_page

            # Perform conversion
            cv.convert(
                str(docx_path),
                start=start,
                end=end,
                multi_processing=self.multi_processing,
                cpu_count=self.cpu_count
            )

            cv.close()

            # Get file stats
            stats = {
                "success": True,
                "input_file": str(pdf_path),
                "output_file": str(docx_path),
                "input_size_mb": pdf_path.stat().st_size / 1024 / 1024,
                "output_size_mb": docx_path.stat().st_size / 1024 / 1024,
                "pages_converted": f"{start}-{end or 'end'}"
            }

            logger.info(
                f"Conversion successful: "
                f"input={stats['input_size_mb']:.2f}MB, "
                f"output={stats['output_size_mb']:.2f}MB"
            )

            return (True, str(docx_path), stats)

        except Exception as e:
            logger.error(f"PDF conversion failed: {e}", exc_info=True)
            return (
                False,
                "",
                {"error": str(e)}
            )

    def convert_and_validate(
        self,
        pdf_path: str,
        docx_path: Optional[str] = None,
        pages: Optional[List[int]] = None
    ) -> Tuple[bool, str, Dict[str, Any]]:
        """
        Convert PDF to DOCX with quality validation.

        Args:
            pdf_path: Path to input PDF file
            docx_path: Path to output DOCX file
            pages: Optional list of specific pages to convert

        Returns:
            Tuple of (success, output_path, validation_results)
        """
        # Perform conversion
        success, output_path, stats = self.convert_pdf_to_docx(
            pdf_path,
            docx_path,
            pages
        )

        if not success:
            return (False, "", stats)

        # Validate conversion quality
        from transit.core.validator import DocumentValidator

        validator = DocumentValidator()

        try:
            from docx import Document
            doc = Document(output_path)

            # Validate document
            issues = validator.validate_document(doc)

            # Add validation results to stats
            stats["validation"] = {
                "issues": issues,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "has_errors": any(severity == "error" for severity, _ in issues),
                "has_warnings": any(severity == "warning" for severity, _ in issues)
            }

            logger.info(
                f"Validation complete: "
                f"paragraphs={stats['validation']['paragraph_count']}, "
                f"tables={stats['validation']['table_count']}, "
                f"issues={len(issues)}"
            )

            return (True, output_path, stats)

        except Exception as e:
            logger.error(f"Validation failed: {e}", exc_info=True)
            stats["validation_error"] = str(e)
            return (True, output_path, stats)

    def batch_convert(
        self,
        pdf_files: List[str],
        output_dir: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Convert multiple PDF files to DOCX.

        Args:
            pdf_files: List of PDF file paths
            output_dir: Optional output directory (default: same as input)

        Returns:
            List of conversion results (dicts with success, paths, stats)
        """
        results = []

        for i, pdf_file in enumerate(pdf_files):
            logger.info(f"Converting {i + 1}/{len(pdf_files)}: {pdf_file}")

            # Determine output path
            if output_dir:
                output_path = Path(output_dir) / Path(pdf_file).with_suffix('.docx').name
            else:
                output_path = None

            # Convert
            success, output, stats = self.convert_and_validate(
                pdf_file,
                str(output_path) if output_path else None
            )

            results.append({
                "input": pdf_file,
                "output": output,
                "success": success,
                "stats": stats
            })

        # Summary
        successful = sum(1 for r in results if r["success"])
        logger.info(
            f"Batch conversion complete: "
            f"{successful}/{len(pdf_files)} successful"
        )

        return results

    def get_pdf_info(self, pdf_path: str) -> Dict[str, Any]:
        """
        Get information about a PDF file.

        Args:
            pdf_path: Path to PDF file

        Returns:
            Dictionary with PDF metadata
        """
        if not self._converter_available:
            return {"error": "pdf2docx not installed"}

        try:
            import PyPDF2

            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                info = {
                    "path": pdf_path,
                    "page_count": len(pdf_reader.pages),
                    "file_size_mb": Path(pdf_path).stat().st_size / 1024 / 1024
                }

                # Try to get metadata
                if pdf_reader.metadata:
                    info["metadata"] = {
                        "title": pdf_reader.metadata.get("/Title"),
                        "author": pdf_reader.metadata.get("/Author"),
                        "subject": pdf_reader.metadata.get("/Subject"),
                        "creator": pdf_reader.metadata.get("/Creator")
                    }

                return info

        except ImportError:
            logger.warning("PyPDF2 not installed, using basic file info")
            return {
                "path": pdf_path,
                "file_size_mb": Path(pdf_path).stat().st_size / 1024 / 1024
            }
        except Exception as e:
            logger.error(f"Failed to get PDF info: {e}")
            return {"error": str(e)}


class PDFPreviewWorkflow:
    """
    Workflow for PDF conversion with user preview and approval.

    Converts PDF to DOCX, allows user to preview, and proceeds with
    translation only if approved.
    """

    def __init__(self, converter: Optional[PDFConverter] = None):
        """
        Initialize preview workflow.

        Args:
            converter: Optional custom PDF converter
        """
        self.converter = converter or PDFConverter()
        logger.info("Initialized PDF preview workflow")

    def interactive_convert(
        self,
        pdf_path: str,
        auto_approve: bool = False
    ) -> Tuple[bool, Optional[str]]:
        """
        Interactive PDF conversion with user approval.

        Args:
            pdf_path: Path to PDF file
            auto_approve: Skip preview and auto-approve (for CLI)

        Returns:
            Tuple of (approved, docx_path)
        """
        # Get PDF info
        pdf_info = self.converter.get_pdf_info(pdf_path)

        if "error" in pdf_info:
            logger.error(f"Cannot get PDF info: {pdf_info['error']}")
            return (False, None)

        logger.info(f"PDF info: {pdf_info}")

        # Convert with validation
        success, docx_path, stats = self.converter.convert_and_validate(pdf_path)

        if not success:
            logger.error(f"Conversion failed: {stats.get('error')}")
            return (False, None)

        logger.info(f"Conversion successful: {docx_path}")

        # Show validation results
        if "validation" in stats:
            validation = stats["validation"]

            logger.info(
                f"Converted document: "
                f"{validation['paragraph_count']} paragraphs, "
                f"{validation['table_count']} tables"
            )

            if validation["issues"]:
                logger.warning(f"Validation found {len(validation['issues'])} issues:")
                for severity, message in validation["issues"]:
                    logger.warning(f"  [{severity.upper()}] {message}")

        # Auto-approve or ask user
        if auto_approve:
            logger.info("Auto-approving conversion (auto_approve=True)")
            return (True, docx_path)

        # In CLI mode, return path for user to review
        logger.info(
            f"Conversion complete. "
            f"Review the file at: {docx_path}"
        )

        return (True, docx_path)

    def convert_for_translation(
        self,
        pdf_path: str,
        keep_docx: bool = True
    ) -> Tuple[bool, Optional[str], Dict[str, Any]]:
        """
        Convert PDF for translation workflow.

        Args:
            pdf_path: Path to PDF file
            keep_docx: Keep intermediate DOCX after translation

        Returns:
            Tuple of (success, docx_path, conversion_info)
        """
        # Convert with validation
        success, docx_path, stats = self.converter.convert_and_validate(pdf_path)

        if not success:
            return (False, None, stats)

        # Prepare conversion info
        info = {
            "original_pdf": pdf_path,
            "converted_docx": docx_path,
            "keep_intermediate": keep_docx,
            "stats": stats
        }

        logger.info(f"PDF ready for translation: {docx_path}")

        return (True, docx_path, info)
