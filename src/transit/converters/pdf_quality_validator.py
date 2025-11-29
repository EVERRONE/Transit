"""Quality validation for PDF to DOCX conversions."""

import logging
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional
from docx import Document

logger = logging.getLogger(__name__)


class PDFConversionQualityValidator:
    """
    Validates quality of PDF to DOCX conversion.

    Checks for common conversion issues like:
    - Missing content
    - Formatting problems
    - Table structure issues
    - Image handling
    """

    def __init__(
        self,
        min_paragraph_count: int = 1,
        check_tables: bool = True,
        check_formatting: bool = True,
        check_images: bool = True
    ):
        """
        Initialize quality validator.

        Args:
            min_paragraph_count: Minimum expected paragraphs
            check_tables: Validate table structure
            check_formatting: Validate text formatting
            check_images: Check for images
        """
        self.min_paragraph_count = min_paragraph_count
        self.check_tables = check_tables
        self.check_formatting = check_formatting
        self.check_images = check_images

        logger.info("Initialized PDF quality validator")

    def validate_conversion(
        self,
        docx_path: str,
        pdf_page_count: Optional[int] = None
    ) -> Tuple[bool, List[Tuple[str, str]], Dict[str, Any]]:
        """
        Validate converted DOCX file.

        Args:
            docx_path: Path to converted DOCX file
            pdf_page_count: Optional original PDF page count for comparison

        Returns:
            Tuple of (is_valid, issues, stats)
        """
        issues = []
        stats = {
            "paragraph_count": 0,
            "table_count": 0,
            "image_count": 0,
            "empty_paragraphs": 0,
            "formatted_paragraphs": 0,
            "has_content": False
        }

        try:
            doc = Document(docx_path)

            # Check basic content
            stats["paragraph_count"] = len(doc.paragraphs)
            stats["table_count"] = len(doc.tables)

            # Validate minimum content
            if stats["paragraph_count"] < self.min_paragraph_count:
                issues.append((
                    "error",
                    f"Too few paragraphs: {stats['paragraph_count']} < {self.min_paragraph_count}"
                ))

            # Check for completely empty document
            non_empty_paragraphs = [
                p for p in doc.paragraphs
                if p.text.strip()
            ]

            stats["empty_paragraphs"] = stats["paragraph_count"] - len(non_empty_paragraphs)

            if not non_empty_paragraphs and stats["table_count"] == 0:
                issues.append((
                    "error",
                    "Document appears to be empty (no text or tables)"
                ))
                stats["has_content"] = False
            else:
                stats["has_content"] = True

            # Check for excessive empty paragraphs
            if stats["paragraph_count"] > 0:
                empty_ratio = stats["empty_paragraphs"] / stats["paragraph_count"]
                if empty_ratio > 0.5:
                    issues.append((
                        "warning",
                        f"High ratio of empty paragraphs: {empty_ratio:.1%}"
                    ))

            # Validate tables
            if self.check_tables and stats["table_count"] > 0:
                table_issues = self._validate_tables(doc)
                issues.extend(table_issues)
                stats["table_validation"] = len(table_issues)

            # Check formatting preservation
            if self.check_formatting:
                formatting_stats = self._check_formatting(doc)
                stats.update(formatting_stats)

                if formatting_stats["formatted_paragraphs"] == 0 and stats["paragraph_count"] > 10:
                    issues.append((
                        "warning",
                        "No formatted text detected - formatting may be lost"
                    ))

            # Check for images
            if self.check_images:
                image_count = self._count_images(doc)
                stats["image_count"] = image_count

                if image_count == 0 and pdf_page_count and pdf_page_count > 1:
                    issues.append((
                        "info",
                        "No images detected in conversion (original may have had images)"
                    ))

            # Compare with PDF page count if provided
            if pdf_page_count:
                stats["pdf_page_count"] = pdf_page_count

                # Heuristic: expect at least some paragraphs per page
                expected_min_paragraphs = pdf_page_count * 2

                if stats["paragraph_count"] < expected_min_paragraphs:
                    issues.append((
                        "warning",
                        f"Low paragraph count ({stats['paragraph_count']}) "
                        f"for {pdf_page_count}-page PDF (expected ~{expected_min_paragraphs}+)"
                    ))

            # Determine overall validity
            has_errors = any(severity == "error" for severity, _ in issues)
            is_valid = stats["has_content"] and not has_errors

            logger.info(
                f"Validation complete: valid={is_valid}, "
                f"issues={len(issues)}, "
                f"paragraphs={stats['paragraph_count']}, "
                f"tables={stats['table_count']}"
            )

            return (is_valid, issues, stats)

        except Exception as e:
            logger.error(f"Validation failed: {e}", exc_info=True)
            issues.append(("error", f"Validation error: {str(e)}"))
            return (False, issues, stats)

    def _validate_tables(self, doc: Document) -> List[Tuple[str, str]]:
        """
        Validate table structure.

        Args:
            doc: Document to validate

        Returns:
            List of issues (severity, message)
        """
        issues = []

        for i, table in enumerate(doc.tables):
            try:
                # Check for empty tables
                if not table.rows:
                    issues.append((
                        "warning",
                        f"Table {i + 1} has no rows"
                    ))
                    continue

                # Check for inconsistent column counts
                col_counts = [len(row.cells) for row in table.rows]

                if len(set(col_counts)) > 1:
                    issues.append((
                        "warning",
                        f"Table {i + 1} has inconsistent column counts: {set(col_counts)}"
                    ))

                # Check for completely empty tables
                has_content = False
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            has_content = True
                            break
                    if has_content:
                        break

                if not has_content:
                    issues.append((
                        "warning",
                        f"Table {i + 1} appears to be empty"
                    ))

            except Exception as e:
                issues.append((
                    "error",
                    f"Error validating table {i + 1}: {str(e)}"
                ))

        return issues

    def _check_formatting(self, doc: Document) -> Dict[str, int]:
        """
        Check text formatting preservation.

        Args:
            doc: Document to check

        Returns:
            Dictionary with formatting stats
        """
        stats = {
            "formatted_paragraphs": 0,
            "bold_runs": 0,
            "italic_runs": 0,
            "underline_runs": 0,
            "colored_runs": 0
        }

        for para in doc.paragraphs:
            has_formatting = False

            for run in para.runs:
                if run.bold:
                    stats["bold_runs"] += 1
                    has_formatting = True

                if run.italic:
                    stats["italic_runs"] += 1
                    has_formatting = True

                if run.underline:
                    stats["underline_runs"] += 1
                    has_formatting = True

                if run.font.color and run.font.color.rgb:
                    stats["colored_runs"] += 1
                    has_formatting = True

            if has_formatting:
                stats["formatted_paragraphs"] += 1

        return stats

    def _count_images(self, doc: Document) -> int:
        """
        Count images in document.

        Args:
            doc: Document to check

        Returns:
            Number of images
        """
        try:
            # Count inline shapes (images)
            image_count = 0

            for para in doc.paragraphs:
                # Check for inline shapes in runs
                for run in para.runs:
                    if hasattr(run, '_element'):
                        # Check for drawing elements
                        drawings = run._element.findall(
                            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
                        )
                        image_count += len(drawings)

            return image_count

        except Exception as e:
            logger.warning(f"Could not count images: {e}")
            return 0

    def generate_report(
        self,
        validation_result: Tuple[bool, List[Tuple[str, str]], Dict[str, Any]]
    ) -> str:
        """
        Generate human-readable validation report.

        Args:
            validation_result: Result from validate_conversion()

        Returns:
            Formatted report string
        """
        is_valid, issues, stats = validation_result

        report_lines = []
        report_lines.append("=" * 60)
        report_lines.append("PDF CONVERSION QUALITY REPORT")
        report_lines.append("=" * 60)
        report_lines.append("")

        # Overall status
        status = "✓ VALID" if is_valid else "✗ INVALID"
        report_lines.append(f"Overall Status: {status}")
        report_lines.append("")

        # Content statistics
        report_lines.append("Content Statistics:")
        report_lines.append(f"  Paragraphs: {stats.get('paragraph_count', 0)}")
        report_lines.append(f"  Tables: {stats.get('table_count', 0)}")
        report_lines.append(f"  Images: {stats.get('image_count', 0)}")
        report_lines.append(f"  Empty paragraphs: {stats.get('empty_paragraphs', 0)}")
        report_lines.append("")

        # Formatting statistics
        if "formatted_paragraphs" in stats:
            report_lines.append("Formatting Statistics:")
            report_lines.append(f"  Formatted paragraphs: {stats['formatted_paragraphs']}")
            report_lines.append(f"  Bold runs: {stats.get('bold_runs', 0)}")
            report_lines.append(f"  Italic runs: {stats.get('italic_runs', 0)}")
            report_lines.append(f"  Underlined runs: {stats.get('underline_runs', 0)}")
            report_lines.append("")

        # Issues
        if issues:
            report_lines.append("Issues Found:")

            errors = [msg for sev, msg in issues if sev == "error"]
            warnings = [msg for sev, msg in issues if sev == "warning"]
            infos = [msg for sev, msg in issues if sev == "info"]

            if errors:
                report_lines.append("  ERRORS:")
                for msg in errors:
                    report_lines.append(f"    - {msg}")

            if warnings:
                report_lines.append("  WARNINGS:")
                for msg in warnings:
                    report_lines.append(f"    - {msg}")

            if infos:
                report_lines.append("  INFO:")
                for msg in infos:
                    report_lines.append(f"    - {msg}")
        else:
            report_lines.append("No issues found.")

        report_lines.append("")
        report_lines.append("=" * 60)

        return "\n".join(report_lines)


class ConversionComparator:
    """
    Compare PDF and DOCX to assess conversion quality.

    Uses heuristics to estimate how much content was preserved.
    """

    def __init__(self):
        """Initialize conversion comparator."""
        logger.info("Initialized conversion comparator")

    def compare_files(
        self,
        pdf_path: str,
        docx_path: str
    ) -> Dict[str, Any]:
        """
        Compare PDF and DOCX files.

        Args:
            pdf_path: Path to original PDF
            docx_path: Path to converted DOCX

        Returns:
            Dictionary with comparison results
        """
        comparison = {
            "pdf_path": pdf_path,
            "docx_path": docx_path,
            "pdf_stats": {},
            "docx_stats": {},
            "quality_score": 0.0,
            "issues": []
        }

        try:
            # Get PDF stats
            pdf_stats = self._get_pdf_stats(pdf_path)
            comparison["pdf_stats"] = pdf_stats

            # Get DOCX stats
            docx_stats = self._get_docx_stats(docx_path)
            comparison["docx_stats"] = docx_stats

            # Calculate quality score
            quality_score = self._calculate_quality_score(pdf_stats, docx_stats)
            comparison["quality_score"] = quality_score

            # Generate issues
            if quality_score < 0.5:
                comparison["issues"].append((
                    "error",
                    f"Low quality score: {quality_score:.1%}"
                ))
            elif quality_score < 0.7:
                comparison["issues"].append((
                    "warning",
                    f"Medium quality score: {quality_score:.1%}"
                ))

            logger.info(f"Comparison complete: quality_score={quality_score:.1%}")

            return comparison

        except Exception as e:
            logger.error(f"Comparison failed: {e}", exc_info=True)
            comparison["error"] = str(e)
            return comparison

    def _get_pdf_stats(self, pdf_path: str) -> Dict[str, Any]:
        """Get PDF file statistics."""
        stats = {"page_count": 0, "file_size_mb": 0}

        try:
            import PyPDF2

            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                stats["page_count"] = len(pdf_reader.pages)

            stats["file_size_mb"] = Path(pdf_path).stat().st_size / 1024 / 1024

        except ImportError:
            logger.warning("PyPDF2 not available, using basic stats")
            stats["file_size_mb"] = Path(pdf_path).stat().st_size / 1024 / 1024
        except Exception as e:
            logger.error(f"Error getting PDF stats: {e}")

        return stats

    def _get_docx_stats(self, docx_path: str) -> Dict[str, Any]:
        """Get DOCX file statistics."""
        stats = {
            "paragraph_count": 0,
            "table_count": 0,
            "word_count": 0,
            "file_size_mb": 0
        }

        try:
            doc = Document(docx_path)

            stats["paragraph_count"] = len(doc.paragraphs)
            stats["table_count"] = len(doc.tables)

            # Count words
            for para in doc.paragraphs:
                stats["word_count"] += len(para.text.split())

            stats["file_size_mb"] = Path(docx_path).stat().st_size / 1024 / 1024

        except Exception as e:
            logger.error(f"Error getting DOCX stats: {e}")

        return stats

    def _calculate_quality_score(
        self,
        pdf_stats: Dict[str, Any],
        docx_stats: Dict[str, Any]
    ) -> float:
        """
        Calculate quality score (0-1).

        Uses heuristics based on:
        - Content presence (paragraphs, tables, words)
        - File size ratio
        - Page/paragraph ratio

        Args:
            pdf_stats: PDF statistics
            docx_stats: DOCX statistics

        Returns:
            Quality score between 0 and 1
        """
        score = 0.0
        checks = 0

        # Check 1: Has content
        if docx_stats.get("paragraph_count", 0) > 0:
            score += 0.3
        checks += 1

        # Check 2: Reasonable paragraph/page ratio
        pdf_pages = pdf_stats.get("page_count", 1)
        docx_paragraphs = docx_stats.get("paragraph_count", 0)

        if pdf_pages > 0:
            para_per_page = docx_paragraphs / pdf_pages

            # Expect at least 2 paragraphs per page on average
            if para_per_page >= 2:
                score += 0.3
            elif para_per_page >= 1:
                score += 0.15

        checks += 1

        # Check 3: Has words
        if docx_stats.get("word_count", 0) > 0:
            score += 0.2
        checks += 1

        # Check 4: File size is reasonable
        pdf_size = pdf_stats.get("file_size_mb", 0)
        docx_size = docx_stats.get("file_size_mb", 0)

        if pdf_size > 0 and docx_size > 0:
            # DOCX should be somewhat close to PDF size (0.5x to 3x)
            ratio = docx_size / pdf_size

            if 0.5 <= ratio <= 3.0:
                score += 0.2
            elif 0.1 <= ratio <= 5.0:
                score += 0.1

        checks += 1

        # Normalize score
        return score if checks == 0 else score

    def generate_comparison_report(self, comparison: Dict[str, Any]) -> str:
        """
        Generate comparison report.

        Args:
            comparison: Result from compare_files()

        Returns:
            Formatted report string
        """
        lines = []
        lines.append("=" * 60)
        lines.append("PDF TO DOCX CONVERSION COMPARISON")
        lines.append("=" * 60)
        lines.append("")

        lines.append(f"PDF File: {comparison['pdf_path']}")
        lines.append(f"DOCX File: {comparison['docx_path']}")
        lines.append("")

        # PDF stats
        pdf_stats = comparison.get("pdf_stats", {})
        lines.append("Original PDF:")
        lines.append(f"  Pages: {pdf_stats.get('page_count', 'unknown')}")
        lines.append(f"  Size: {pdf_stats.get('file_size_mb', 0):.2f} MB")
        lines.append("")

        # DOCX stats
        docx_stats = comparison.get("docx_stats", {})
        lines.append("Converted DOCX:")
        lines.append(f"  Paragraphs: {docx_stats.get('paragraph_count', 0)}")
        lines.append(f"  Tables: {docx_stats.get('table_count', 0)}")
        lines.append(f"  Words: {docx_stats.get('word_count', 0)}")
        lines.append(f"  Size: {docx_stats.get('file_size_mb', 0):.2f} MB")
        lines.append("")

        # Quality score
        quality = comparison.get("quality_score", 0)
        lines.append(f"Quality Score: {quality:.1%}")
        lines.append("")

        # Issues
        if comparison.get("issues"):
            lines.append("Issues:")
            for severity, msg in comparison["issues"]:
                lines.append(f"  [{severity.upper()}] {msg}")
        else:
            lines.append("No issues detected.")

        lines.append("")
        lines.append("=" * 60)

        return "\n".join(lines)
