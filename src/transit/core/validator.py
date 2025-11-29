"""Document validation utilities."""

from docx import Document
from typing import List, Tuple
import logging

logger = logging.getLogger(__name__)


class DocumentValidator:
    """Validate documents before and after translation."""

    @staticmethod
    def validate_document(doc: Document) -> List[Tuple[str, str]]:
        """
        Validate document before processing.

        Args:
            doc: Document to validate

        Returns:
            List of (severity, message) tuples
        """
        issues = []

        # Check for tables
        if doc.tables:
            issues.append(("info", f"Document contains {len(doc.tables)} table(s)"))

        # Check for sections
        if len(doc.sections) > 1:
            issues.append(("info", f"Document has {len(doc.sections)} sections"))

        # Check paragraph count
        para_count = len(doc.paragraphs)
        if para_count > 1000:
            issues.append(("warning", f"Large document: {para_count} paragraphs. Processing may take time."))

        return issues

    @staticmethod
    def validate_translation_output(
        original_path: str,
        translated_doc: Document
    ) -> List[Tuple[str, str]]:
        """
        Verify translation preserved structure.

        Args:
            original_path: Path to original document
            translated_doc: Translated document

        Returns:
            List of (severity, message) tuples
        """
        checks = []

        try:
            original_doc = Document(original_path)

            # Paragraph count should approximately double (original + translation)
            # Note: Not exact because of empty paragraphs and edge cases
            orig_para_count = len(list(original_doc.paragraphs))
            trans_para_count = len(list(translated_doc.paragraphs))

            # Allow 10% variance
            expected_min = orig_para_count * 1.5
            expected_max = orig_para_count * 2.5

            if not (expected_min <= trans_para_count <= expected_max):
                checks.append((
                    "warning",
                    f"Paragraph count unusual: {orig_para_count} → {trans_para_count}"
                ))

            # Table count should remain same
            orig_table_count = len(original_doc.tables)
            trans_table_count = len(translated_doc.tables)

            if orig_table_count != trans_table_count:
                checks.append((
                    "error",
                    f"Table count changed: {orig_table_count} → {trans_table_count}"
                ))

            # Section count should remain same
            if len(original_doc.sections) != len(translated_doc.sections):
                checks.append(("error", "Section count mismatch"))

        except Exception as e:
            checks.append(("error", f"Validation failed: {e}"))

        return checks
