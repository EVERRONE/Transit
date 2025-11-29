"""Core document processor with run-level translation."""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from typing import Iterable, List, Dict, Any, Optional
import logging
from tqdm import tqdm

from transit.utils.formatting import clone_run_formatting, clone_paragraph_formatting
from transit.utils.list_formatting import preserve_list_structure_in_translation
from transit.utils.hyperlink_formatting import preserve_hyperlinks_in_translation
from transit.core.exceptions import CorruptDocumentError
from transit.utils import docx_patch as _docx_patch  # noqa: F401
from transit.parsers.context_collection import (
    collect_document_contexts,
    collect_section_contexts,
    ParagraphContext,
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process DOCX documents with run-level translation."""

    def __init__(self, translator):
        """
        Initialize document processor.

        Args:
            translator: Translator instance (expected to provide OpenAI-style interface)
        """
        self.translator = translator
        self.supports_context = hasattr(translator, 'set_document_context')

    def translate_document(
        self,
        input_path: str,
        output_path: str,
        target_lang: str,
        show_progress: bool = False
    ) -> None:
        """
        Main translation pipeline.

        Args:
            input_path: Path to input DOCX file
            output_path: Path to output DOCX file
            target_lang: Target language code (e.g., "EN-US")
            show_progress: Show progress bar

        Raises:
            CorruptDocumentError: If document structure is invalid
        """
        try:
            doc = Document(input_path)
            logger.info(f"Loaded document: {input_path}")
        except Exception as e:
            raise CorruptDocumentError(f"Cannot load document: {e}")

        # Set document context for OpenAI (helps with abbreviations, technical terms)
        if self.supports_context:
            context = self._extract_document_context(doc)
            self.translator.set_document_context(context)
            logger.info("Document context set for intelligent translation")

        # Collect every paragraph (body + headers/footers + nested tables)
        traversal = collect_document_contexts(doc)
        contexts = traversal.contexts

        if traversal.warnings:
            for warning in traversal.warnings:
                logger.warning(warning)

        logger.info("Translating %d paragraphs (including headers/footers)...", len(contexts))
        iterator: Iterable[ParagraphContext]
        if show_progress:
            iterator = tqdm(contexts, desc="Translating", total=len(contexts))
        else:
            iterator = contexts

        for context in iterator:
            paragraph = context.paragraph
            self._translate_paragraph(paragraph, target_lang)

        # Save output
        try:
            doc.save(output_path)
            logger.info(f"Saved translated document: {output_path}")
        except Exception as e:
            raise CorruptDocumentError(f"Cannot save document: {e}")

    def _translate_paragraph(self, paragraph: Paragraph, target_lang: str) -> None:
        """
        Translate single paragraph at run+sentence level.

        Args:
            paragraph: Paragraph to translate
            target_lang: Target language code
        """
        # Skip empty paragraphs
        if not paragraph.text.strip():
            return

        # For OpenAI: translate entire paragraph for better context
        self._translate_paragraph_openai(paragraph, target_lang)

    def _translate_paragraph_openai(self, paragraph: Paragraph, target_lang: str) -> None:
        """
        Translate paragraph using OpenAI (context-aware, paragraph-level).

        Args:
            paragraph: Paragraph to translate
            target_lang: Target language code
        """
        # Get full paragraph text for context-aware translation
        full_text = paragraph.text

        translated_full = self.translator.translate_text(
            full_text,
            target_lang=target_lang,
            source_lang="NL",
            preserve_formatting=True
        )

        self._apply_translated_text(paragraph, translated_full)

    def _apply_translated_text(self, paragraph: Paragraph, translated_text: str) -> None:
        """
        Insert translated text underneath the paragraph while preserving formatting.

        Args:
            paragraph: Paragraph that has been translated
            translated_text: Translated paragraph text
        """
        if not translated_text:
            return

        original_text = paragraph.text

        translated_runs: List[Dict[str, Any]] = []
        current_pos = 0
        total_length = len(original_text)

        for run in paragraph.runs:
            text = run.text or ""
            if not text:
                continue

            run_length = len(text)

            if total_length > 0:
                progress = (current_pos + run_length) / total_length
                translated_limit = int(progress * len(translated_text))

                if translated_runs:
                    prev_end = sum(len(entry['text']) for entry in translated_runs)
                else:
                    prev_end = 0

                translated_chunk = translated_text[prev_end:translated_limit]
            else:
                translated_chunk = ""

            if translated_chunk:
                translated_runs.append({
                    "text": translated_chunk,
                    "original_run": run,
                })

            current_pos += run_length

        if translated_runs:
            total_translated = sum(len(entry['text']) for entry in translated_runs)
            if total_translated < len(translated_text):
                translated_runs[-1]['text'] += translated_text[total_translated:]
        elif paragraph.runs:
            translated_runs.append({
                "text": translated_text,
                "original_run": paragraph.runs[0],
            })

        if translated_runs:
            self._insert_translation_paragraph_after(paragraph, translated_runs)

    def _insert_translation_paragraph_after(
        self,
        original_paragraph: Paragraph,
        translated_runs: List[Dict[str, Any]]
    ) -> Paragraph:
        """
        Insert translation paragraph directly after original.
        Uses low-level XML manipulation.

        Args:
            original_paragraph: Original paragraph
            translated_runs: List of dicts with 'text' and 'original_run'

        Returns:
            Created translation paragraph
        """
        # Create new paragraph via XML
        new_p = parse_xml(r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')

        # Insert in document order via XML
        original_paragraph._p.addnext(new_p)

        # Wrap in python-docx Paragraph object
        translation_para = Paragraph(new_p, original_paragraph._parent)

        # Clone paragraph formatting
        clone_paragraph_formatting(original_paragraph, translation_para)

        # Preserve list formatting if applicable (bullets, numbering)
        preserve_list_structure_in_translation(original_paragraph, translation_para)

        # Preserve hyperlinks (logs hyperlinks for awareness)
        preserve_hyperlinks_in_translation(original_paragraph, translation_para)

        # Add translated runs with formatting
        for run_data in translated_runs:
            new_run = translation_para.add_run(run_data['text'])
            clone_run_formatting(run_data['original_run'], new_run)
            # Visual marker voor vertaling
            new_run.italic = True

        return translation_para

    def _translate_table(self, table: Table, target_lang: str) -> None:
        """
        Translate table cell-by-cell, including nested tables.

        Args:
            table: Table to translate
            target_lang: Target language code
        """
        processed_cells = set()
        row_count = len(table.rows)
        col_count = len(table.columns) if table.rows else 0

        for row_idx in range(row_count):
            for col_idx in range(col_count):
                cell = table.cell(row_idx, col_idx)
                cell_id = id(cell._element)

                # Skip if already processed (merged cell)
                if cell_id in processed_cells:
                    continue

                processed_cells.add(cell_id)

                # Process nested tables first (if any)
                nested_tables = cell.tables
                if nested_tables:
                    logger.info(f"Found {len(nested_tables)} nested table(s) in cell ({row_idx}, {col_idx})")
                    for nested_table in nested_tables:
                        self._translate_table(nested_table, target_lang)  # Recursive

                # Process each paragraph in cell
                for paragraph in cell.paragraphs:
                    self._translate_paragraph(paragraph, target_lang)

    def _translate_section_headers_footers(self, section, target_lang: str) -> None:
        """
        Translate headers and footers in section.

        Args:
            section: Document section
            target_lang: Target language code
        """
        entries = [
            ("header", section.header),
            ("header_first_page", section.first_page_header),
            ("header_even_page", section.even_page_header),
            ("footer", section.footer),
            ("footer_first_page", section.first_page_footer),
            ("footer_even_page", section.even_page_footer),
        ]

        translated = 0
        for location, header_footer in entries:
            contexts = collect_section_contexts(header_footer, location=location)
            for context in contexts:
                self._translate_paragraph(context.paragraph, target_lang)
                translated += 1

        if translated:
            logger.info("Translated %d header/footer paragraphs for section", translated)

    def _extract_document_context(self, doc: Document) -> str:
        """
        Extract document context for intelligent translation.

        This helps OpenAI understand abbreviations, domain, and style.

        Args:
            doc: Document to extract context from

        Returns:
            Context string describing document
        """
        # Collect sample paragraphs (first 5 non-empty)
        sample_texts = []
        paragraphs_attr = getattr(doc, 'paragraphs', [])
        try:
            paragraphs_iterable = list(paragraphs_attr)
        except TypeError:
            paragraphs_iterable = []

        for paragraph in paragraphs_iterable[:10]:
            text = getattr(paragraph, 'text', '')
            text = text.strip() if isinstance(text, str) else ''
            if text and len(text) > 10:
                sample_texts.append(text)
                if len(sample_texts) >= 5:
                    break

        # Extract headers for context
        header_texts = []
        sections_attr = getattr(doc, 'sections', [])
        try:
            sections_iterable = list(sections_attr)
        except TypeError:
            sections_iterable = []

        for section in sections_iterable:
            header = getattr(section, 'header', None)
            if header and not getattr(header, 'is_linked_to_previous', False):
                for para in getattr(header, 'paragraphs', []):
                    text = getattr(para, 'text', '')
                    text = text.strip() if isinstance(text, str) else ''
                    if text:
                        header_texts.append(text)
                        break  # Just first header paragraph

        # Build context summary
        context_parts = []

        if header_texts:
            context_parts.append(f"Document headers: {' | '.join(header_texts[:3])}")

        if sample_texts:
            context_parts.append(f"Sample content:\n" + "\n".join(sample_texts[:3]))

        # Detect document type based on content
        full_text = " ".join(sample_texts).lower()
        if any(word in full_text for word in ["artikel", "sectie", "paragraaf", "wet", "verordening"]):
            context_parts.append("Document type: Legal/regulatory document")
        elif any(word in full_text for word in ["rapport", "analyse", "onderzoek", "conclusie"]):
            context_parts.append("Document type: Report/analysis")
        elif any(word in full_text for word in ["contract", "overeenkomst", "partijen"]):
            context_parts.append("Document type: Contract/agreement")
        elif any(word in full_text for word in ["instructie", "handleiding", "stap", "procedure"]):
            context_parts.append("Document type: Instructions/manual")

        return "\n".join(context_parts) if context_parts else "General document"
