"""Preview panel for document content."""

import tkinter as tk
from tkinter import ttk, scrolledtext
from typing import Optional
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class PreviewPanel(ttk.LabelFrame):
    """
    Panel for previewing document content.

    Shows basic information about the document before/after translation.
    """

    def __init__(self, parent, **kwargs):
        """
        Initialize preview panel.

        Args:
            parent: Parent widget
            **kwargs: Additional frame arguments
        """
        super().__init__(parent, text="Document Preview", padding="10", **kwargs)

        self.current_file: Optional[str] = None

        self._create_widgets()
        self._setup_layout()

    def _create_widgets(self):
        """Create preview widgets."""
        # Info frame
        self.info_frame = ttk.Frame(self)

        self.filename_label = ttk.Label(
            self.info_frame,
            text="No document loaded",
            font=('Helvetica', 10, 'bold')
        )

        self.stats_label = ttk.Label(
            self.info_frame,
            text="",
            font=('Helvetica', 9)
        )

        # Preview text area
        self.preview_text = scrolledtext.ScrolledText(
            self,
            height=15,
            width=80,
            wrap='word',
            state='disabled',
            font=('Courier', 9)
        )

        # Refresh button
        self.refresh_button = ttk.Button(
            self,
            text="Refresh Preview",
            command=self._refresh_preview
        )

    def _setup_layout(self):
        """Layout widgets."""
        self.info_frame.pack(fill='x', pady=(0, 10))
        self.filename_label.pack(anchor='w')
        self.stats_label.pack(anchor='w', pady=(5, 0))

        self.preview_text.pack(fill='both', expand=True, pady=(0, 10))
        self.refresh_button.pack()

    def load_file(self, file_path: str):
        """
        Load and preview file.

        Args:
            file_path: Path to file to preview
        """
        self.current_file = file_path
        self._refresh_preview()

    def _refresh_preview(self):
        """Refresh preview content."""
        if not self.current_file:
            return

        try:
            file_path = Path(self.current_file)

            # Update filename
            self.filename_label.config(text=f"File: {file_path.name}")

            # Get file stats
            stats_text = self._get_file_stats(file_path)
            self.stats_label.config(text=stats_text)

            # Get preview content
            preview_content = self._get_preview_content(file_path)

            # Update preview text
            self.preview_text.config(state='normal')
            self.preview_text.delete('1.0', 'end')
            self.preview_text.insert('1.0', preview_content)
            self.preview_text.config(state='disabled')

            logger.info(f"Preview loaded for: {file_path}")

        except Exception as e:
            logger.error(f"Error loading preview: {e}", exc_info=True)
            self.preview_text.config(state='normal')
            self.preview_text.delete('1.0', 'end')
            self.preview_text.insert('1.0', f"Error loading preview:\n{str(e)}")
            self.preview_text.config(state='disabled')

    def _get_file_stats(self, file_path: Path) -> str:
        """
        Get file statistics.

        Args:
            file_path: Path to file

        Returns:
            Statistics string
        """
        try:
            file_size_mb = file_path.stat().st_size / 1024 / 1024

            if file_path.suffix.lower() == '.pdf':
                return self._get_pdf_stats(file_path, file_size_mb)
            elif file_path.suffix.lower() == '.docx':
                return self._get_docx_stats(file_path, file_size_mb)
            else:
                return f"Size: {file_size_mb:.2f} MB"

        except Exception as e:
            logger.error(f"Error getting file stats: {e}")
            return "Unable to load statistics"

    def _get_pdf_stats(self, file_path: Path, file_size_mb: float) -> str:
        """Get PDF file statistics."""
        try:
            import PyPDF2

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                page_count = len(pdf_reader.pages)

            return f"Type: PDF | Pages: {page_count} | Size: {file_size_mb:.2f} MB"

        except ImportError:
            return f"Type: PDF | Size: {file_size_mb:.2f} MB (install PyPDF2 for more details)"
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return f"Type: PDF | Size: {file_size_mb:.2f} MB"

    def _get_docx_stats(self, file_path: Path, file_size_mb: float) -> str:
        """Get DOCX file statistics."""
        try:
            from docx import Document

            doc = Document(file_path)

            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # Count words
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)

            return (
                f"Type: DOCX | Paragraphs: {paragraph_count} | "
                f"Tables: {table_count} | Words: {word_count} | "
                f"Size: {file_size_mb:.2f} MB"
            )

        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return f"Type: DOCX | Size: {file_size_mb:.2f} MB"

    def _get_preview_content(self, file_path: Path) -> str:
        """
        Get preview content from file.

        Args:
            file_path: Path to file

        Returns:
            Preview text
        """
        if file_path.suffix.lower() == '.pdf':
            return self._get_pdf_preview(file_path)
        elif file_path.suffix.lower() == '.docx':
            return self._get_docx_preview(file_path)
        else:
            return "Preview not available for this file type."

    def _get_pdf_preview(self, file_path: Path) -> str:
        """Get preview of PDF content."""
        try:
            import PyPDF2

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                # Get text from first few pages
                preview_lines = ["PDF CONTENT PREVIEW", "=" * 60, ""]

                max_pages = min(3, len(pdf_reader.pages))

                for i in range(max_pages):
                    page = pdf_reader.pages[i]
                    text = page.extract_text()

                    preview_lines.append(f"--- Page {i + 1} ---")
                    preview_lines.append(text[:500] + "..." if len(text) > 500 else text)
                    preview_lines.append("")

                if len(pdf_reader.pages) > max_pages:
                    preview_lines.append(f"... and {len(pdf_reader.pages) - max_pages} more pages")

                return "\n".join(preview_lines)

        except ImportError:
            return (
                "PDF preview not available.\n\n"
                "Install PyPDF2 for PDF preview support:\n"
                "pip install PyPDF2"
            )
        except Exception as e:
            logger.error(f"Error previewing PDF: {e}")
            return f"Error previewing PDF:\n{str(e)}"

    def _get_docx_preview(self, file_path: Path) -> str:
        """Get preview of DOCX content."""
        try:
            from docx import Document

            doc = Document(file_path)

            preview_lines = ["DOCX CONTENT PREVIEW", "=" * 60, ""]

            # Get first few paragraphs
            max_paragraphs = 20
            paragraph_count = 0

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                preview_lines.append(para.text)
                paragraph_count += 1

                if paragraph_count >= max_paragraphs:
                    break

            if len(doc.paragraphs) > paragraph_count:
                preview_lines.append("")
                preview_lines.append(f"... and {len(doc.paragraphs) - paragraph_count} more paragraphs")

            # Show tables info
            if doc.tables:
                preview_lines.append("")
                preview_lines.append(f"Document contains {len(doc.tables)} table(s)")

            return "\n".join(preview_lines)

        except Exception as e:
            logger.error(f"Error previewing DOCX: {e}")
            return f"Error previewing DOCX:\n{str(e)}"

    def clear(self):
        """Clear preview."""
        self.current_file = None
        self.filename_label.config(text="No document loaded")
        self.stats_label.config(text="")

        self.preview_text.config(state='normal')
        self.preview_text.delete('1.0', 'end')
        self.preview_text.config(state='disabled')
