"""Performance benchmarks for TransIt."""

import time
import os
from docx import Document
from unittest.mock import Mock
from transit.parsers.document_processor import DocumentProcessor
from transit.translators.openai_translator import OpenAITranslator
import statistics


class MockTranslator:
    """Mock translator that simulates API delays."""

    def __init__(self, delay_ms=50):
        """
        Initialize mock translator.

        Args:
            delay_ms: Simulated API delay in milliseconds
        """
        self.delay_ms = delay_ms
        self.call_count = 0
        self.total_chars = 0

    def set_document_context(self, context: str):
        """Set document context."""
        pass

    def translate_text(self, text: str, target_lang: str, source_lang: str = "NL",
                      preserve_formatting: bool = True, context: str = None) -> str:
        """Mock translation with simulated delay."""
        if not text or not text.strip():
            return text

        # Simulate API delay
        time.sleep(self.delay_ms / 1000.0)

        self.call_count += 1
        self.total_chars += len(text)

        return text.upper()


class Benchmark:
    """Performance benchmark suite."""

    def __init__(self):
        self.results = []

    def run_benchmark(self, name: str, func, iterations=3):
        """
        Run a benchmark function multiple times and collect stats.

        Args:
            name: Benchmark name
            func: Function to benchmark
            iterations: Number of iterations
        """
        print(f"\n{'='*60}")
        print(f"Benchmark: {name}")
        print(f"{'='*60}")

        times = []
        for i in range(iterations):
            print(f"  Iteration {i+1}/{iterations}...", end=" ")
            start = time.time()
            result = func()
            elapsed = time.time() - start
            times.append(elapsed)
            print(f"{elapsed:.3f}s")

        avg_time = statistics.mean(times)
        min_time = min(times)
        max_time = max(times)
        std_dev = statistics.stdev(times) if len(times) > 1 else 0

        print(f"\n  Results:")
        print(f"    Average: {avg_time:.3f}s")
        print(f"    Min:     {min_time:.3f}s")
        print(f"    Max:     {max_time:.3f}s")
        print(f"    StdDev:  {std_dev:.3f}s")

        if result:
            print(f"    Details: {result}")

        self.results.append({
            'name': name,
            'avg_time': avg_time,
            'min_time': min_time,
            'max_time': max_time,
            'std_dev': std_dev,
            'result': result
        })

    def print_summary(self):
        """Print summary of all benchmarks."""
        print(f"\n{'='*60}")
        print("BENCHMARK SUMMARY")
        print(f"{'='*60}")
        print(f"{'Benchmark':<40} {'Avg Time':>12}")
        print(f"{'-'*40} {'-'*12}")

        for result in self.results:
            print(f"{result['name']:<40} {result['avg_time']:>10.3f}s")


def benchmark_simple_document(tmp_path):
    """Benchmark: Simple document (10 paragraphs)."""
    doc_path = tmp_path / "simple_10.docx"
    output_path = tmp_path / "output_simple_10.docx"

    # Create test document
    doc = Document()
    for i in range(10):
        doc.add_paragraph(f"Dit is paragraaf nummer {i+1}. Het bevat wat tekst om te vertalen.")
    doc.save(str(doc_path))

    # Benchmark
    translator = MockTranslator(delay_ms=10)
    processor = DocumentProcessor(translator)

    processor.translate_document(str(doc_path), str(output_path), "EN-US")

    return {
        'paragraphs': 10,
        'api_calls': translator.call_count,
        'total_chars': translator.total_chars
    }


def benchmark_medium_document(tmp_path):
    """Benchmark: Medium document (50 paragraphs)."""
    doc_path = tmp_path / "medium_50.docx"
    output_path = tmp_path / "output_medium_50.docx"

    # Create test document
    doc = Document()
    for i in range(50):
        doc.add_paragraph(f"Paragraaf {i+1}. " + "Dit is een standaard zin. " * 3)
    doc.save(str(doc_path))

    # Benchmark
    translator = MockTranslator(delay_ms=10)
    processor = DocumentProcessor(translator)

    processor.translate_document(str(doc_path), str(output_path), "EN-US")

    return {
        'paragraphs': 50,
        'api_calls': translator.call_count,
        'total_chars': translator.total_chars
    }


def benchmark_large_document(tmp_path):
    """Benchmark: Large document (200 paragraphs)."""
    doc_path = tmp_path / "large_200.docx"
    output_path = tmp_path / "output_large_200.docx"

    # Create test document
    doc = Document()
    for i in range(200):
        doc.add_paragraph(f"Paragraaf {i+1}. " + "Tekst. " * 5)
    doc.save(str(doc_path))

    # Benchmark
    translator = MockTranslator(delay_ms=10)
    processor = DocumentProcessor(translator)

    processor.translate_document(str(doc_path), str(output_path), "EN-US")

    return {
        'paragraphs': 200,
        'api_calls': translator.call_count,
        'total_chars': translator.total_chars
    }


def benchmark_table_document(tmp_path):
    """Benchmark: Document with tables (10 tables, 3x3 each)."""
    doc_path = tmp_path / "tables_10.docx"
    output_path = tmp_path / "output_tables_10.docx"

    # Create test document
    doc = Document()
    for t in range(10):
        doc.add_paragraph(f"Tabel {t+1}")
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table.cell(i, j).text = f"Cel ({i},{j}) in tabel {t+1}"
    doc.save(str(doc_path))

    # Benchmark
    translator = MockTranslator(delay_ms=10)
    processor = DocumentProcessor(translator)

    processor.translate_document(str(doc_path), str(output_path), "EN-US")

    return {
        'tables': 10,
        'cells': 90,
        'api_calls': translator.call_count,
        'total_chars': translator.total_chars
    }


def benchmark_mixed_document(tmp_path):
    """Benchmark: Mixed document (paragraphs + tables + formatting)."""
    doc_path = tmp_path / "mixed.docx"
    output_path = tmp_path / "output_mixed.docx"

    # Create test document
    doc = Document()

    # Add header
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Document Header"

    # Add paragraphs
    for i in range(20):
        doc.add_paragraph(f"Paragraaf {i+1}. Dit is normale tekst.")

    # Add tables
    for t in range(3):
        table = doc.add_table(rows=4, cols=3)
        for i in range(4):
            for j in range(3):
                table.cell(i, j).text = f"Data {i},{j}"

    # Add more paragraphs
    for i in range(10):
        para = doc.add_paragraph()
        para.add_run("Tekst met ").bold = False
        run = para.add_run("vet")
        run.bold = True
        para.add_run(" en ").bold = False
        run2 = para.add_run("schuin")
        run2.italic = True

    # Add footer
    footer = section.footer
    footer.paragraphs[0].text = "Document Footer"

    doc.save(str(doc_path))

    # Benchmark
    translator = MockTranslator(delay_ms=10)
    processor = DocumentProcessor(translator)

    processor.translate_document(str(doc_path), str(output_path), "EN-US")

    return {
        'paragraphs': 30,
        'tables': 3,
        'api_calls': translator.call_count,
        'total_chars': translator.total_chars
    }


def benchmark_api_delay_impact(tmp_path):
    """Benchmark: Impact of API delay on performance."""
    doc_path = tmp_path / "delay_test.docx"
    output_path = tmp_path / "output_delay.docx"

    # Create test document
    doc = Document()
    for i in range(20):
        doc.add_paragraph(f"Paragraaf {i+1}.")
    doc.save(str(doc_path))

    results = []

    for delay_ms in [10, 50, 100, 200]:
        print(f"    Testing with {delay_ms}ms delay...", end=" ")
        translator = MockTranslator(delay_ms=delay_ms)
        processor = DocumentProcessor(translator)

        start = time.time()
        processor.translate_document(str(doc_path), str(output_path), "EN-US")
        elapsed = time.time() - start

        results.append({
            'delay_ms': delay_ms,
            'time': elapsed,
            'calls': translator.call_count
        })
        print(f"{elapsed:.3f}s ({translator.call_count} calls)")

    return results


def main():
    """Run all benchmarks."""
    import tempfile
    from pathlib import Path

    print("\n" + "="*60)
    print("TransIt Performance Benchmarks")
    print("="*60)

    # Create temporary directory for test files
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)

        bench = Benchmark()

        # Run benchmarks
        bench.run_benchmark(
            "Simple Document (10 paragraphs)",
            lambda: benchmark_simple_document(tmp_path),
            iterations=3
        )

        bench.run_benchmark(
            "Medium Document (50 paragraphs)",
            lambda: benchmark_medium_document(tmp_path),
            iterations=3
        )

        bench.run_benchmark(
            "Large Document (200 paragraphs)",
            lambda: benchmark_large_document(tmp_path),
            iterations=2  # Fewer iterations for large docs
        )

        bench.run_benchmark(
            "Table Document (10 tables, 3x3)",
            lambda: benchmark_table_document(tmp_path),
            iterations=3
        )

        bench.run_benchmark(
            "Mixed Document (paras + tables + format)",
            lambda: benchmark_mixed_document(tmp_path),
            iterations=3
        )

        bench.run_benchmark(
            "API Delay Impact Test",
            lambda: benchmark_api_delay_impact(tmp_path),
            iterations=1  # Only once, tests multiple delays internally
        )

        # Print summary
        bench.print_summary()

        print("\n" + "="*60)
        print("Benchmarks Complete")
        print("="*60)
        print("\nNotes:")
        print("- Mock translator simulates 10ms API delay per call")
        print("- Real OpenAI delays vary (50-500ms typical)")
        print("- Actual performance depends on:")
        print("  * API latency and rate limits")
        print("  * Document complexity")
        print("  * Network conditions")
        print("  * System resources")


if __name__ == "__main__":
    main()
