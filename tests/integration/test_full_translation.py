"""Integration tests for full document translation pipeline."""

import pytest
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from unittest.mock import Mock, patch
from transit.parsers.document_processor import DocumentProcessor
from transit.translators.openai_translator import OpenAITranslator


class MockOpenAITranslator:
    """Mock OpenAI translator for testing without API calls."""

    def __init__(self):
        self.document_context = None

    def set_document_context(self, context: str):
        """Set document context."""
        self.document_context = context

    def translate_text(self, text: str, target_lang: str, source_lang: str = "NL",
                      preserve_formatting: bool = True, context: str = None) -> str:
        """Mock translation - returns uppercase version."""
        if not text or not text.strip():
            return text
        return text.upper()

    def translate_batch(self, texts: list, target_lang: str, source_lang: str = "NL",
                       preserve_formatting: bool = True, batch_context: str = None) -> list:
        """Mock batch translation."""
        return [t.upper() if t and t.strip() else t for t in texts]


@pytest.fixture
def mock_translator():
    """Create mock translator."""
    return MockOpenAITranslator()


@pytest.fixture
def temp_docx_path(tmp_path):
    """Create temporary path for DOCX files."""
    return tmp_path / "test.docx"


@pytest.fixture
def temp_output_path(tmp_path):
    """Create temporary path for output files."""
    return tmp_path / "output.docx"


class TestSimpleDocumentTranslation:
    """Test translation of simple documents."""

    def test_single_paragraph_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating document with single paragraph."""
        # Create test document
        doc = Document()
        doc.add_paragraph("Dit is een test.")
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        paras = list(output_doc.paragraphs)

        # Should have 2 paragraphs (original + translation)
        assert len(paras) == 2
        assert paras[0].text == "Dit is een test."
        assert paras[1].text == "DIT IS EEN TEST."

    def test_multiple_paragraphs_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating document with multiple paragraphs."""
        # Create test document
        doc = Document()
        doc.add_paragraph("Eerste paragraaf.")
        doc.add_paragraph("Tweede paragraaf.")
        doc.add_paragraph("Derde paragraaf.")
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        paras = list(output_doc.paragraphs)

        # Should have 6 paragraphs (3 original + 3 translations)
        assert len(paras) == 6
        assert paras[0].text == "Eerste paragraaf."
        assert paras[1].text == "EERSTE PARAGRAAF."
        assert paras[2].text == "Tweede paragraaf."
        assert paras[3].text == "TWEEDE PARAGRAAF."


class TestFormattingPreservation:
    """Test that formatting is preserved during translation."""

    def test_bold_preservation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test that bold formatting is preserved."""
        # Create test document with bold text
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Vetgedrukte tekst")
        run.bold = True
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        translated_para = output_doc.paragraphs[1]

        # Translation should also be bold (but italic as marker)
        assert translated_para.runs[0].bold is True
        assert translated_para.runs[0].italic is True  # Visual marker

    def test_italic_preservation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test that italic formatting is preserved."""
        # Create test document with italic text
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Schuingedrukte tekst")
        run.italic = True
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        translated_para = output_doc.paragraphs[1]

        # Translation should be italic
        assert translated_para.runs[0].italic is True

    def test_font_size_preservation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test that font size is preserved."""
        # Create test document with custom font size
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Grote tekst")
        run.font.size = Pt(24)
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        translated_para = output_doc.paragraphs[1]

        # Translation should have same font size
        assert translated_para.runs[0].font.size == Pt(24)

    def test_mixed_formatting_preservation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test that mixed formatting in paragraph is preserved."""
        # Create test document with mixed formatting
        doc = Document()
        para = doc.add_paragraph()
        run1 = para.add_run("Normale tekst ")
        run2 = para.add_run("vetgedrukt")
        run2.bold = True
        run3 = para.add_run(" en normaal")
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        translated_para = output_doc.paragraphs[1]

        # Should have 3 runs with same formatting pattern
        assert len(translated_para.runs) == 3
        assert translated_para.runs[0].bold is not True  # First run not bold
        assert translated_para.runs[1].bold is True       # Second run bold
        assert translated_para.runs[2].bold is not True  # Third run not bold

    def test_paragraph_alignment_preservation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test that paragraph alignment is preserved."""
        # Create test document with centered paragraph
        doc = Document()
        para = doc.add_paragraph("Gecentreerde tekst")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        translated_para = output_doc.paragraphs[1]

        # Translation should also be centered
        assert translated_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


class TestTableTranslation:
    """Test translation of tables."""

    def test_simple_table_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating a simple 2x2 table."""
        # Create test document with table
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cel 1"
        table.cell(0, 1).text = "Cel 2"
        table.cell(1, 0).text = "Cel 3"
        table.cell(1, 1).text = "Cel 4"
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        output_table = output_doc.tables[0]

        # Table structure should be preserved
        assert len(output_table.rows) == 2
        assert len(output_table.columns) == 2

        # Each cell should have original + translation
        cell_00_paras = output_table.cell(0, 0).paragraphs
        assert len(cell_00_paras) == 2
        assert cell_00_paras[0].text == "Cel 1"
        assert cell_00_paras[1].text == "CEL 1"

    def test_table_with_empty_cells(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating table with empty cells."""
        # Create test document with table containing empty cells
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Tekst"
        # Leave other cells empty
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output - should not crash
        output_doc = Document(str(temp_output_path))
        assert len(output_doc.tables) == 1


class TestHeaderFooterTranslation:
    """Test translation of headers and footers."""

    def test_header_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating document header."""
        # Create test document with header
        doc = Document()
        doc.add_paragraph("Body text")

        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "Dit is een header"
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        output_header = output_doc.sections[0].header

        # Header should have original + translation
        header_paras = output_header.paragraphs
        assert len(header_paras) >= 2
        # Note: may have extra empty paragraphs, so check content
        texts = [p.text for p in header_paras if p.text.strip()]
        assert "Dit is een header" in texts
        assert "DIT IS EEN HEADER" in texts

    def test_footer_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating document footer."""
        # Create test document with footer
        doc = Document()
        doc.add_paragraph("Body text")

        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Dit is een footer"
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        output_footer = output_doc.sections[0].footer

        # Footer should have original + translation
        footer_paras = output_footer.paragraphs
        texts = [p.text for p in footer_paras if p.text.strip()]
        assert "Dit is een footer" in texts
        assert "DIT IS EEN FOOTER" in texts


class TestComplexDocuments:
    """Test translation of complex documents."""

    def test_document_with_multiple_elements(self, mock_translator, temp_docx_path, temp_output_path):
        """Test document with paragraphs, tables, headers, and footers."""
        # Create complex test document
        doc = Document()

        # Add header
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "Header tekst"

        # Add paragraphs
        doc.add_paragraph("Eerste paragraaf")

        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Tabel cel"

        # Add more paragraphs
        doc.add_paragraph("Tweede paragraaf")

        # Add footer
        footer = section.footer
        footer.paragraphs[0].text = "Footer tekst"

        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output exists and is valid
        output_doc = Document(str(temp_output_path))

        # Should have paragraphs
        assert len(output_doc.paragraphs) >= 4

        # Should have table
        assert len(output_doc.tables) == 1

        # Should have header and footer
        assert len(output_doc.sections[0].header.paragraphs) >= 1
        assert len(output_doc.sections[0].footer.paragraphs) >= 1


class TestListFormatting:
    """Test translation of lists with formatting preservation."""

    def test_bullet_list_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating bullet list preserves list formatting."""
        # Create test document with bullet list
        doc = Document()
        doc.add_paragraph("Eerste item", style='List Bullet')
        doc.add_paragraph("Tweede item", style='List Bullet')
        doc.add_paragraph("Derde item", style='List Bullet')
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        paras = list(output_doc.paragraphs)

        # Should have 6 paragraphs (3 original + 3 translations)
        assert len(paras) == 6

        # Check that translations are also formatted as lists
        from transit.utils.list_formatting import has_list_formatting

        # Original items should be lists
        assert has_list_formatting(paras[0]) is True
        assert has_list_formatting(paras[2]) is True
        assert has_list_formatting(paras[4]) is True

        # Translation items should also be lists
        assert has_list_formatting(paras[1]) is True
        assert has_list_formatting(paras[3]) is True
        assert has_list_formatting(paras[5]) is True

    def test_numbered_list_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating numbered list preserves numbering."""
        # Create test document with numbered list
        doc = Document()
        doc.add_paragraph("Eerste stap", style='List Number')
        doc.add_paragraph("Tweede stap", style='List Number')
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        paras = list(output_doc.paragraphs)

        from transit.utils.list_formatting import has_list_formatting

        # All should be lists (originals and translations)
        assert len(paras) == 4
        for para in paras:
            assert has_list_formatting(para) is True

    def test_nested_list_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating nested list preserves indentation levels."""
        # Create test document with nested list
        doc = Document()
        doc.add_paragraph("Hoofditem", style='List Bullet')
        doc.add_paragraph("Sub-item", style='List Bullet 2')
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        paras = list(output_doc.paragraphs)

        from transit.utils.list_formatting import get_list_level

        # Should have 4 paragraphs
        assert len(paras) == 4

        # Check indentation levels are preserved
        # Original level 0 and its translation
        assert get_list_level(paras[0]) == 0
        assert get_list_level(paras[1]) == 0

        # Original level 1 and its translation
        assert get_list_level(paras[2]) > 0
        assert get_list_level(paras[3]) > 0

        # Translation should have same level as original
        assert get_list_level(paras[2]) == get_list_level(paras[3])

    def test_mixed_list_and_paragraph(self, mock_translator, temp_docx_path, temp_output_path):
        """Test document with mix of lists and regular paragraphs."""
        # Create test document
        doc = Document()
        doc.add_paragraph("Normale paragraaf")
        doc.add_paragraph("Lijst item 1", style='List Bullet')
        doc.add_paragraph("Lijst item 2", style='List Bullet')
        doc.add_paragraph("Nog een normale paragraaf")
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output
        output_doc = Document(str(temp_output_path))
        paras = list(output_doc.paragraphs)

        from transit.utils.list_formatting import has_list_formatting

        # Should have 8 paragraphs (4 original + 4 translations)
        assert len(paras) == 8

        # Regular paragraphs (original and translation)
        assert has_list_formatting(paras[0]) is False
        assert has_list_formatting(paras[1]) is False

        # List items (original and translation)
        assert has_list_formatting(paras[2]) is True
        assert has_list_formatting(paras[3]) is True
        assert has_list_formatting(paras[4]) is True
        assert has_list_formatting(paras[5]) is True

        # Regular paragraphs again
        assert has_list_formatting(paras[6]) is False
        assert has_list_formatting(paras[7]) is False


class TestNestedTables:
    """Test translation of nested tables."""

    def test_nested_table_translation(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating document with nested tables."""
        # Create test document with nested table
        doc = Document()

        # Main table
        main_table = doc.add_table(rows=2, cols=2)
        main_table.cell(0, 0).text = "Hoofd tabel cel"

        # Nested table in cell (0, 1)
        cell_with_nested = main_table.cell(0, 1)
        nested_table = cell_with_nested.add_table(rows=2, cols=2)
        nested_table.cell(0, 0).text = "Geneste cel 1"
        nested_table.cell(0, 1).text = "Geneste cel 2"

        main_table.cell(1, 0).text = "Nog een cel"
        main_table.cell(1, 1).text = "Laatste cel"

        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output exists and is valid
        output_doc = Document(str(temp_output_path))

        # Should have main table
        assert len(output_doc.tables) == 1

        # Main table should have nested table
        output_main_table = output_doc.tables[0]
        output_cell_with_nested = output_main_table.cell(0, 1)

        # Check for nested table
        assert len(output_cell_with_nested.tables) == 1

        output_nested_table = output_cell_with_nested.tables[0]

        # Nested table should have paragraphs translated
        nested_cell_paras = output_nested_table.cell(0, 0).paragraphs

        # Should have original + translation
        assert len(nested_cell_paras) >= 2

    def test_deeply_nested_tables(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating deeply nested tables (3 levels)."""
        # Create test document
        doc = Document()

        # Level 1: Main table
        level1_table = doc.add_table(rows=1, cols=1)
        level1_table.cell(0, 0).text = "Niveau 1"

        # Level 2: Nested table
        cell_level1 = level1_table.cell(0, 0)
        level2_table = cell_level1.add_table(rows=1, cols=1)
        level2_table.cell(0, 0).text = "Niveau 2"

        # Level 3: Deeply nested table
        cell_level2 = level2_table.cell(0, 0)
        level3_table = cell_level2.add_table(rows=1, cols=1)
        level3_table.cell(0, 0).text = "Niveau 3"

        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Should not crash - deep nesting handled
        output_doc = Document(str(temp_output_path))
        assert len(output_doc.tables) >= 1


class TestEdgeCases:
    """Test edge cases in full pipeline."""

    def test_empty_document(self, mock_translator, temp_docx_path, temp_output_path):
        """Test translating empty document."""
        # Create empty test document
        doc = Document()
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Verify output exists
        assert os.path.exists(temp_output_path)
        output_doc = Document(str(temp_output_path))
        # May have one empty paragraph by default
        assert len(output_doc.paragraphs) >= 0

    def test_document_with_only_whitespace(self, mock_translator, temp_docx_path, temp_output_path):
        """Test document with only whitespace paragraphs."""
        # Create test document with whitespace
        doc = Document()
        doc.add_paragraph("   ")
        doc.add_paragraph("\t\n")
        doc.save(str(temp_docx_path))

        # Process
        processor = DocumentProcessor(mock_translator)
        processor.translate_document(
            str(temp_docx_path),
            str(temp_output_path),
            "EN-US"
        )

        # Should not crash
        assert os.path.exists(temp_output_path)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
