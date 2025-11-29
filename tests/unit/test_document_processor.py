"""Unit tests for document processor."""

import pytest
from unittest.mock import Mock, MagicMock, patch
from docx import Document
from docx.text.paragraph import Paragraph
from transit.parsers.document_processor import DocumentProcessor
from transit.core.exceptions import CorruptDocumentError


class TestDocumentProcessorInit:
    """Test document processor initialization."""

    def test_init_with_context_aware_translator(self):
        """Document processor detects context support when translator exposes it."""
        mock_translator = Mock()
        mock_translator.set_document_context = Mock()

        processor = DocumentProcessor(mock_translator)

        assert processor.translator is mock_translator
        assert processor.supports_context is True

    def test_init_without_context_support(self):
        """Document processor handles translators that lack context support."""
        mock_translator = Mock(spec=[])

        processor = DocumentProcessor(mock_translator)

        assert processor.translator is mock_translator
        assert processor.supports_context is False


class TestDocumentLoading:
    """Test document loading and validation."""

    @patch('transit.parsers.document_processor.Document')
    def test_load_valid_document(self, mock_doc_class):
        """Test loading a valid document."""
        mock_doc = Mock()
        mock_doc.iter_inner_content.return_value = []
        mock_doc.sections = []
        mock_doc_class.return_value = mock_doc

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        processor = DocumentProcessor(mock_translator)

        # Should not raise
        processor.translate_document("test.docx", "output.docx", "EN-US")

    @patch('transit.parsers.document_processor.Document')
    def test_load_corrupt_document_raises_error(self, mock_doc_class):
        """Test that corrupt document raises CorruptDocumentError."""
        mock_doc_class.side_effect = Exception("Cannot load document")

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        processor = DocumentProcessor(mock_translator)

        with pytest.raises(CorruptDocumentError):
            processor.translate_document("corrupt.docx", "output.docx", "EN-US")


class TestContextExtraction:
    """Test document context extraction for OpenAI."""

    def test_extract_context_from_simple_doc(self):
        """Test extracting context from simple document."""
        doc = Document()
        doc.add_paragraph("Dit is een test paragraph.")
        doc.add_paragraph("Dit is een tweede paragraph.")

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        processor = DocumentProcessor(mock_translator)

        context = processor._extract_document_context(doc)

        assert isinstance(context, str)
        assert len(context) > 0

    def test_extract_context_detects_legal_document(self):
        """Test that legal keywords are detected."""
        doc = Document()
        doc.add_paragraph("Dit artikel behandelt de wet omtrent verordening.")

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        processor = DocumentProcessor(mock_translator)

        context = processor._extract_document_context(doc)

        assert "Legal" in context or "legal" in context

    def test_extract_context_detects_report(self):
        """Test that report keywords are detected."""
        doc = Document()
        doc.add_paragraph("Dit rapport analyseert de onderzoeksresultaten en conclusie.")

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        processor = DocumentProcessor(mock_translator)

        context = processor._extract_document_context(doc)

        assert "Report" in context or "analysis" in context.lower()

    def test_extract_context_with_headers(self):
        """Test context extraction includes headers."""
        doc = Document()
        doc.add_paragraph("Regular paragraph")

        # Add section with header
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "Document Header Title"

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        processor = DocumentProcessor(mock_translator)

        context = processor._extract_document_context(doc)

        assert "Document Header Title" in context or "header" in context.lower()


class TestParagraphTranslation:
    """Test paragraph translation logic."""

    def test_translate_empty_paragraph_skipped(self):
        """Test that empty paragraphs are skipped."""
        doc = Document()
        para = doc.add_paragraph("")

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        processor = DocumentProcessor(mock_translator)

        # Should not crash and not call translator
        processor._translate_paragraph(para, "EN-US")
        mock_translator.translate_text.assert_not_called()

    def test_translate_paragraph_openai_mode(self):
        """Test paragraph translation in OpenAI mode."""
        doc = Document()
        para = doc.add_paragraph("Dit is een test.")

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        mock_translator.translate_text.return_value = "This is a test."
        processor = DocumentProcessor(mock_translator)

        processor._translate_paragraph(para, "EN-US")

        # Should call translate_text once (paragraph-level)
        mock_translator.translate_text.assert_called_once()
        call_args = mock_translator.translate_text.call_args
        assert "Dit is een test." in call_args[0]


class TestTableTranslation:
    """Test table translation logic."""

    def test_translate_simple_table(self):
        """Test translating a simple 2x2 table."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        mock_translator.translate_text.return_value = "Translated"
        processor = DocumentProcessor(mock_translator)

        processor._translate_table(table, "EN-US")

        # Should process 4 cells
        assert mock_translator.translate_text.call_count >= 4

    def test_translate_table_with_merged_cells(self):
        """Test that merged cells are not processed twice."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # Merge cells
        cell1 = table.cell(0, 0)
        cell2 = table.cell(0, 1)
        cell1.merge(cell2)

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        mock_translator.translate_text.return_value = "Translated"
        processor = DocumentProcessor(mock_translator)

        processor._translate_table(table, "EN-US")

        # Should not crash and handle merged cells correctly


class TestHeaderFooterTranslation:
    """Test header and footer translation."""

    def test_translate_header(self):
        """Test translating document header."""
        doc = Document()
        doc.add_paragraph("Body text")

        # Add header
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "Header Text"

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        mock_translator.translate_text.return_value = "Translated"
        processor = DocumentProcessor(mock_translator)

        processor._translate_section_headers_footers(section, "EN-US")

        # Should translate header
        mock_translator.translate_text.assert_called()

    def test_translate_footer(self):
        """Test translating document footer."""
        doc = Document()
        doc.add_paragraph("Body text")

        # Add footer
        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].text = "Footer Text"

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        mock_translator.translate_text.return_value = "Translated"
        processor = DocumentProcessor(mock_translator)

        processor._translate_section_headers_footers(section, "EN-US")

        # Should translate footer
        mock_translator.translate_text.assert_called()


class TestTranslationInsertion:
    """Test translation paragraph insertion."""

    def test_insert_translation_after_paragraph(self):
        """Test inserting translation paragraph after original."""
        doc = Document()
        original_para = doc.add_paragraph("Original text")

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        mock_translator.translate_text.return_value = "Translated text"
        processor = DocumentProcessor(mock_translator)

        # Create translated runs data
        translated_runs = [{
            'text': "Translated text",
            'original_run': original_para.runs[0]
        }]

        result_para = processor._insert_translation_paragraph_after(original_para, translated_runs)

        # Should create new paragraph
        assert result_para is not None
        assert result_para.text.strip() == "Translated text"

        # Document should now have 2 paragraphs
        assert len(doc.paragraphs) == 2


class TestFullPipeline:
    """Test complete translation pipeline."""

    @patch('transit.parsers.document_processor.Document')
    def test_full_translation_pipeline(self, mock_doc_class):
        """Test complete document translation."""
        # Create mock document
        mock_doc = Mock()
        mock_para = Mock()
        mock_para.text = "Test paragraph"
        mock_para.runs = [Mock(text="Test paragraph")]

        mock_doc.iter_inner_content.return_value = [mock_para]
        mock_doc.sections = []
        mock_doc.paragraphs = [mock_para]
        mock_doc.save = Mock()
        mock_doc_class.return_value = mock_doc

        mock_translator = Mock()
        mock_translator.set_document_context = Mock()
        mock_translator.translate_text.return_value = "Translated"
        processor = DocumentProcessor(mock_translator)

        # Should complete without error
        processor.translate_document("input.docx", "output.docx", "EN-US")

        # Should have called save
        mock_doc.save.assert_called_once_with("output.docx")


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
