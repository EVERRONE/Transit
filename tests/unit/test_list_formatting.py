"""Unit tests for list formatting preservation."""

import pytest
from docx import Document
from docx.oxml import parse_xml
from transit.utils.list_formatting import (
    has_list_formatting,
    get_list_properties,
    clone_list_formatting,
    is_bullet_list,
    get_list_level,
    preserve_list_structure_in_translation
)


@pytest.fixture
def doc_with_lists():
    """Create document with various list types."""
    doc = Document()

    # Bulleted list
    doc.add_paragraph("Bulleted item 1", style='List Bullet')
    doc.add_paragraph("Bulleted item 2", style='List Bullet')

    # Numbered list
    doc.add_paragraph("Numbered item 1", style='List Number')
    doc.add_paragraph("Numbered item 2", style='List Number')

    # Nested list
    doc.add_paragraph("Level 0 item", style='List Bullet')
    doc.add_paragraph("Level 1 item", style='List Bullet 2')

    # Regular paragraph (no list)
    doc.add_paragraph("Regular paragraph")

    return doc


class TestHasListFormatting:
    """Test detection of list formatting."""

    def test_bullet_list_detected(self, doc_with_lists):
        """Test that bullet list is detected."""
        para = doc_with_lists.paragraphs[0]  # "Bulleted item 1"
        assert has_list_formatting(para) is True

    def test_numbered_list_detected(self, doc_with_lists):
        """Test that numbered list is detected."""
        para = doc_with_lists.paragraphs[2]  # "Numbered item 1"
        assert has_list_formatting(para) is True

    def test_nested_list_detected(self, doc_with_lists):
        """Test that nested list is detected."""
        para = doc_with_lists.paragraphs[4]  # "Level 0 item"
        assert has_list_formatting(para) is True

    def test_regular_paragraph_not_detected(self, doc_with_lists):
        """Test that regular paragraph is not detected as list."""
        para = doc_with_lists.paragraphs[6]  # "Regular paragraph"
        assert has_list_formatting(para) is False

    def test_empty_paragraph_not_detected(self):
        """Test that empty paragraph is not detected as list."""
        doc = Document()
        para = doc.add_paragraph("")
        assert has_list_formatting(para) is False


class TestGetListProperties:
    """Test extraction of list properties."""

    def test_get_bullet_list_properties(self, doc_with_lists):
        """Test getting properties from bullet list."""
        para = doc_with_lists.paragraphs[0]
        props = get_list_properties(para)

        # Should have numId and ilvl
        assert 'numId' in props
        assert 'ilvl' in props

        # Top level should be ilvl 0
        assert props['ilvl'] == '0'

    def test_get_numbered_list_properties(self, doc_with_lists):
        """Test getting properties from numbered list."""
        para = doc_with_lists.paragraphs[2]
        props = get_list_properties(para)

        assert 'numId' in props
        assert 'ilvl' in props

    def test_nested_list_has_higher_ilvl(self, doc_with_lists):
        """Test that nested list has higher ilvl."""
        para_level0 = doc_with_lists.paragraphs[4]  # Level 0
        para_level1 = doc_with_lists.paragraphs[5]  # Level 1

        props_level0 = get_list_properties(para_level0)
        props_level1 = get_list_properties(para_level1)

        # Level 1 should have higher ilvl than level 0
        ilvl0 = int(props_level0.get('ilvl', '0'))
        ilvl1 = int(props_level1.get('ilvl', '0'))

        assert ilvl1 > ilvl0

    def test_regular_paragraph_has_no_properties(self, doc_with_lists):
        """Test that regular paragraph returns empty dict."""
        para = doc_with_lists.paragraphs[6]
        props = get_list_properties(para)

        assert props == {}


class TestCloneListFormatting:
    """Test cloning of list formatting."""

    def test_clone_bullet_list_formatting(self, doc_with_lists):
        """Test cloning bullet list formatting to new paragraph."""
        source_para = doc_with_lists.paragraphs[0]  # Bullet list
        target_para = doc_with_lists.add_paragraph("New paragraph")

        # Initially target should not be a list
        assert has_list_formatting(target_para) is False

        # Clone formatting
        clone_list_formatting(source_para, target_para)

        # Now target should be a list
        assert has_list_formatting(target_para) is True

        # Properties should match
        source_props = get_list_properties(source_para)
        target_props = get_list_properties(target_para)

        assert source_props['numId'] == target_props['numId']
        assert source_props['ilvl'] == target_props['ilvl']

    def test_clone_numbered_list_formatting(self, doc_with_lists):
        """Test cloning numbered list formatting."""
        source_para = doc_with_lists.paragraphs[2]  # Numbered list
        target_para = doc_with_lists.add_paragraph("New paragraph")

        clone_list_formatting(source_para, target_para)

        assert has_list_formatting(target_para) is True

        source_props = get_list_properties(source_para)
        target_props = get_list_properties(target_para)

        assert source_props['numId'] == target_props['numId']

    def test_clone_from_regular_paragraph_does_nothing(self, doc_with_lists):
        """Test that cloning from regular paragraph does nothing."""
        source_para = doc_with_lists.paragraphs[6]  # Regular paragraph
        target_para = doc_with_lists.add_paragraph("New paragraph")

        # Both should not be lists
        assert has_list_formatting(source_para) is False
        assert has_list_formatting(target_para) is False

        # Clone should do nothing
        clone_list_formatting(source_para, target_para)

        # Target should still not be a list
        assert has_list_formatting(target_para) is False

    def test_clone_preserves_indentation_level(self, doc_with_lists):
        """Test that cloning preserves indentation level."""
        source_para = doc_with_lists.paragraphs[5]  # Nested list (level 1)
        target_para = doc_with_lists.add_paragraph("New paragraph")

        clone_list_formatting(source_para, target_para)

        source_level = get_list_level(source_para)
        target_level = get_list_level(target_para)

        assert source_level == target_level
        assert target_level > 0  # Should be nested


class TestGetListLevel:
    """Test getting list indentation level."""

    def test_top_level_is_zero(self, doc_with_lists):
        """Test that top-level list has level 0."""
        para = doc_with_lists.paragraphs[0]  # Top-level bullet
        level = get_list_level(para)

        assert level == 0

    def test_nested_level_is_greater(self, doc_with_lists):
        """Test that nested list has level > 0."""
        para = doc_with_lists.paragraphs[5]  # Nested list
        level = get_list_level(para)

        assert level > 0

    def test_regular_paragraph_is_zero(self, doc_with_lists):
        """Test that regular paragraph returns 0."""
        para = doc_with_lists.paragraphs[6]
        level = get_list_level(para)

        assert level == 0


class TestPreserveListStructureInTranslation:
    """Test main preservation function."""

    def test_preserve_bullet_list(self, doc_with_lists):
        """Test preserving bullet list structure."""
        original_para = doc_with_lists.paragraphs[0]  # Bullet list
        translation_para = doc_with_lists.add_paragraph("Translated text")

        # Apply preservation
        preserve_list_structure_in_translation(original_para, translation_para)

        # Translation should now be a list
        assert has_list_formatting(translation_para) is True

        # Properties should match
        orig_props = get_list_properties(original_para)
        trans_props = get_list_properties(translation_para)

        assert orig_props['numId'] == trans_props['numId']
        assert orig_props['ilvl'] == trans_props['ilvl']

    def test_preserve_numbered_list(self, doc_with_lists):
        """Test preserving numbered list structure."""
        original_para = doc_with_lists.paragraphs[2]  # Numbered list
        translation_para = doc_with_lists.add_paragraph("Vertaalde tekst")

        preserve_list_structure_in_translation(original_para, translation_para)

        assert has_list_formatting(translation_para) is True

    def test_preserve_nested_list_level(self, doc_with_lists):
        """Test preserving nested list level."""
        original_para = doc_with_lists.paragraphs[5]  # Nested list
        translation_para = doc_with_lists.add_paragraph("Translation")

        preserve_list_structure_in_translation(original_para, translation_para)

        orig_level = get_list_level(original_para)
        trans_level = get_list_level(translation_para)

        assert orig_level == trans_level

    def test_regular_paragraph_unchanged(self, doc_with_lists):
        """Test that regular paragraph is not affected."""
        original_para = doc_with_lists.paragraphs[6]  # Regular paragraph
        translation_para = doc_with_lists.add_paragraph("Translation")

        # Neither should be lists
        assert has_list_formatting(original_para) is False
        assert has_list_formatting(translation_para) is False

        # Preserve should do nothing
        preserve_list_structure_in_translation(original_para, translation_para)

        # Translation should still not be a list
        assert has_list_formatting(translation_para) is False


class TestEdgeCases:
    """Test edge cases in list formatting."""

    def test_empty_list_item(self):
        """Test handling empty list item."""
        doc = Document()
        para = doc.add_paragraph("", style='List Bullet')

        # Should still be detected as list
        assert has_list_formatting(para) is True

    def test_whitespace_only_list_item(self):
        """Test handling whitespace-only list item."""
        doc = Document()
        para = doc.add_paragraph("   ", style='List Bullet')

        assert has_list_formatting(para) is True

    def test_clone_to_existing_list_item(self, doc_with_lists):
        """Test cloning to paragraph that's already a list."""
        source_para = doc_with_lists.paragraphs[0]  # Bullet list
        target_para = doc_with_lists.paragraphs[2]  # Numbered list (different)

        # Get original properties
        orig_target_props = get_list_properties(target_para)

        # Clone should replace existing list formatting
        clone_list_formatting(source_para, target_para)

        # Should now have source properties
        new_target_props = get_list_properties(target_para)
        source_props = get_list_properties(source_para)

        assert new_target_props['numId'] == source_props['numId']


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
