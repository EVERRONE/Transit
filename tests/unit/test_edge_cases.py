"""Unit tests for edge cases in document processing."""

import pytest
from docx import Document
from docx.shared import Pt, RGBColor


class TestEmptyContent:
    """Test handling of empty or whitespace-only content."""

    def test_empty_paragraph(self):
        """Empty paragraphs should be skipped."""
        doc = Document()
        para = doc.add_paragraph("")

        # Empty paragraph should have no text
        assert para.text == ""
        assert len(para.runs) == 0

    def test_whitespace_only_paragraph(self):
        """Whitespace-only paragraphs should be handled."""
        doc = Document()
        para = doc.add_paragraph("   \t\n   ")

        # Should have text but it's all whitespace
        assert para.text.strip() == ""

    def test_empty_run_in_paragraph(self):
        """Paragraphs with empty runs should be handled."""
        doc = Document()
        para = doc.add_paragraph()
        run1 = para.add_run("")
        run2 = para.add_run("Text")
        run3 = para.add_run("")

        # Should have 3 runs but only middle one has text
        assert len(para.runs) == 3
        assert para.runs[1].text == "Text"


class TestSpecialCharacters:
    """Test handling of special characters."""

    def test_non_breaking_space(self):
        """Non-breaking spaces should be preserved."""
        doc = Document()
        text_with_nbsp = "Test\u00A0text"  # \u00A0 = non-breaking space
        para = doc.add_paragraph(text_with_nbsp)

        assert "\u00A0" in para.text

    def test_tab_character(self):
        """Tab characters should be preserved."""
        doc = Document()
        text_with_tab = "Test\ttext"
        para = doc.add_paragraph(text_with_tab)

        assert "\t" in para.text

    def test_line_break(self):
        """Line breaks within paragraph should be preserved."""
        doc = Document()
        para = doc.add_paragraph("Line 1")
        para.add_run("\n")
        para.add_run("Line 2")

        assert "\n" in para.text

    def test_multiple_spaces(self):
        """Multiple consecutive spaces should be preserved."""
        doc = Document()
        text = "Test    multiple    spaces"
        para = doc.add_paragraph(text)

        assert "    " in para.text



class TestFormattingEdgeCases:
    """Test edge cases in formatting preservation."""

    def test_mixed_bold_italic(self):
        """Run with both bold and italic."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold and italic")
        run.bold = True
        run.italic = True

        assert run.bold is True
        assert run.italic is True

    def test_font_size_edge_cases(self):
        """Very small and very large font sizes."""
        doc = Document()
        para = doc.add_paragraph()

        # Very small
        run1 = para.add_run("Small")
        run1.font.size = Pt(6)
        assert run1.font.size == Pt(6)

        # Very large
        run2 = para.add_run("Large")
        run2.font.size = Pt(72)
        assert run2.font.size == Pt(72)

    def test_color_edge_cases(self):
        """Black, white, and custom colors."""
        doc = Document()
        para = doc.add_paragraph()

        # Black
        run1 = para.add_run("Black")
        run1.font.color.rgb = RGBColor(0, 0, 0)

        # White
        run2 = para.add_run("White")
        run2.font.color.rgb = RGBColor(255, 255, 255)

        # Custom
        run3 = para.add_run("Purple")
        run3.font.color.rgb = RGBColor(128, 0, 128)

    def test_tri_state_none(self):
        """Tri-state properties set to None (inherited)."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Inherited")

        # By default, bold/italic should be None (inherited)
        assert run.bold is None
        assert run.italic is None
        assert run.underline is None

    def test_superscript_subscript_mutually_exclusive(self):
        """Superscript and subscript are mutually exclusive."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Text")

        run.font.superscript = True
        assert run.font.superscript is True
        assert run.font.subscript is not True

        run.font.subscript = True
        assert run.font.subscript is True
        assert run.font.superscript is not True


class TestTableEdgeCases:
    """Test edge cases in table processing."""

    def test_empty_table_cell(self):
        """Empty table cells should be handled."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # Leave one cell empty
        table.cell(0, 0).text = "Text"
        # table.cell(0, 1) is empty

        assert table.cell(0, 0).text == "Text"
        assert table.cell(0, 1).text == ""

    def test_table_with_single_cell(self):
        """1x1 table should work."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Single cell"

        assert len(table.rows) == 1
        assert len(table.columns) == 1

    def test_table_with_many_rows(self):
        """Table with many rows should be handled."""
        doc = Document()
        table = doc.add_table(rows=100, cols=2)

        assert len(table.rows) == 100

    def test_merged_cells_horizontal(self):
        """Horizontally merged cells."""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)

        # Merge cells in first row
        cell1 = table.cell(0, 0)
        cell2 = table.cell(0, 2)
        cell1.merge(cell2)

        # After merge, accessing same cell from different positions
        # should return same cell
        assert table.cell(0, 0) is table.cell(0, 1)

    def test_merged_cells_vertical(self):
        """Vertically merged cells."""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)

        # Merge cells in first column
        cell1 = table.cell(0, 0)
        cell2 = table.cell(2, 0)
        cell1.merge(cell2)

        assert table.cell(0, 0) is table.cell(1, 0)


class TestParagraphEdgeCases:
    """Test edge cases in paragraph processing."""

    def test_paragraph_with_no_runs(self):
        """Paragraph with no runs (just created)."""
        doc = Document()
        para = doc.add_paragraph()

        assert len(para.runs) == 0
        assert para.text == ""

    def test_paragraph_with_many_runs(self):
        """Paragraph with many runs."""
        doc = Document()
        para = doc.add_paragraph()

        for i in range(50):
            para.add_run(f"Run {i} ")

        assert len(para.runs) == 50

    def test_paragraph_only_formatting_no_text(self):
        """Paragraph with runs that have formatting but no text."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("")
        run.bold = True
        run.font.size = Pt(12)

        assert run.text == ""
        assert run.bold is True

    def test_very_long_paragraph(self):
        """Very long paragraph with lots of text."""
        doc = Document()
        long_text = "A" * 10000  # 10k characters
        para = doc.add_paragraph(long_text)

        assert len(para.text) == 10000


