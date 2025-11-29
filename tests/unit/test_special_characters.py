"""Unit tests for special character handling."""

import pytest
from docx import Document
from transit.utils.special_characters import (
    protect_special_characters,
    restore_special_characters,
    has_tabs,
    get_tab_count,
    add_tab,
    preserve_tabs_in_run,
    has_line_break,
    get_line_break_count,
    add_line_break,
    preserve_line_breaks_in_run,
    detect_special_whitespace,
    normalize_whitespace,
    preserve_special_formatting_in_run,
    SPECIAL_CHARS,
    SPECIAL_CHARS_REVERSE
)


class TestProtectRestore:
    """Test protection and restoration of special characters."""

    def test_protect_non_breaking_space(self):
        """Test protecting non-breaking space."""
        text = "text\u00A0with\u00A0nbsp"
        protected = protect_special_characters(text)

        assert '\u00A0' not in protected
        assert '<NBSP>' in protected

    def test_protect_tab(self):
        """Test protecting tab character."""
        text = "text\twith\ttab"
        protected = protect_special_characters(text)

        assert '\t' not in protected
        assert '<TAB>' in protected

    def test_protect_multiple_special_chars(self):
        """Test protecting multiple types of special characters."""
        text = "text\twith\u00A0different\tspecial\u00A0chars"
        protected = protect_special_characters(text)

        assert '\t' not in protected
        assert '\u00A0' not in protected
        assert '<TAB>' in protected
        assert '<NBSP>' in protected

    def test_restore_non_breaking_space(self):
        """Test restoring non-breaking space."""
        text = "text<NBSP>with<NBSP>nbsp"
        restored = restore_special_characters(text)

        assert '<NBSP>' not in restored
        assert '\u00A0' in restored

    def test_restore_tab(self):
        """Test restoring tab character."""
        text = "text<TAB>with<TAB>tab"
        restored = restore_special_characters(text)

        assert '<TAB>' not in restored
        assert '\t' in restored

    def test_protect_and_restore_roundtrip(self):
        """Test protect and restore cycle preserves text."""
        original = "text\twith\u00A0special\tchars\u00A0here"

        protected = protect_special_characters(original)
        restored = restore_special_characters(protected)

        assert restored == original

    def test_protect_empty_text(self):
        """Test protecting empty text."""
        assert protect_special_characters("") == ""
        assert protect_special_characters(None) is None

    def test_restore_empty_text(self):
        """Test restoring empty text."""
        assert restore_special_characters("") == ""
        assert restore_special_characters(None) is None

    def test_protect_text_without_special_chars(self):
        """Test protecting text without special characters."""
        text = "regular text with spaces"
        protected = protect_special_characters(text)

        assert protected == text


class TestTabHandling:
    """Test tab detection and preservation."""

    def test_has_tabs_with_tab_element(self):
        """Test detecting tab element in run."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("text")

        # Initially no tabs
        assert has_tabs(run) is False

        # Add tab
        add_tab(run)

        # Now should have tab
        assert has_tabs(run) is True

    def test_get_tab_count(self):
        """Test counting tabs in run."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("text")

        # Initially 0 tabs
        assert get_tab_count(run) == 0

        # Add one tab
        add_tab(run)
        assert get_tab_count(run) == 1

        # Add another
        add_tab(run)
        assert get_tab_count(run) == 2

    def test_add_tab(self):
        """Test adding tab to run."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("text")

        add_tab(run)

        assert has_tabs(run) is True
        assert get_tab_count(run) == 1

    def test_preserve_tabs_in_run(self):
        """Test preserving tabs from original to translation."""
        doc = Document()
        para = doc.add_paragraph()

        original_run = para.add_run("original")
        add_tab(original_run)
        add_tab(original_run)

        translation_run = para.add_run("translation")

        # Initially translation has no tabs
        assert get_tab_count(translation_run) == 0

        # Preserve tabs
        preserve_tabs_in_run(original_run, translation_run)

        # Now should have same number
        assert get_tab_count(translation_run) == 2


class TestLineBreakHandling:
    """Test line break detection and preservation."""

    def test_has_line_break(self):
        """Test detecting line break in run."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("text")

        # Initially no line breaks
        assert has_line_break(run) is False

        # Add line break
        add_line_break(run)

        # Now should have line break
        assert has_line_break(run) is True

    def test_get_line_break_count(self):
        """Test counting line breaks in run."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("text")

        assert get_line_break_count(run) == 0

        add_line_break(run)
        assert get_line_break_count(run) == 1

        add_line_break(run)
        assert get_line_break_count(run) == 2

    def test_add_line_break(self):
        """Test adding line break to run."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("text")

        add_line_break(run)

        assert has_line_break(run) is True

    def test_preserve_line_breaks_in_run(self):
        """Test preserving line breaks from original to translation."""
        doc = Document()
        para = doc.add_paragraph()

        original_run = para.add_run("original")
        add_line_break(original_run)

        translation_run = para.add_run("translation")

        assert get_line_break_count(translation_run) == 0

        preserve_line_breaks_in_run(original_run, translation_run)

        assert get_line_break_count(translation_run) == 1


class TestDetectSpecialWhitespace:
    """Test detection of special whitespace."""

    def test_detect_non_breaking_space(self):
        """Test detecting non-breaking spaces."""
        text = "text\u00A0with\u00A0nbsp"
        counts = detect_special_whitespace(text)

        assert '<NBSP>' in counts
        assert counts['<NBSP>'] == 2

    def test_detect_tab(self):
        """Test detecting tabs."""
        text = "text\twith\ttab"
        counts = detect_special_whitespace(text)

        assert 'REGULAR_TAB' in counts
        assert counts['REGULAR_TAB'] == 2

    def test_detect_multiple_types(self):
        """Test detecting multiple special character types."""
        text = "text\twith\u00A0different\tspecial\u00A0chars"
        counts = detect_special_whitespace(text)

        assert 'REGULAR_TAB' in counts
        assert '<NBSP>' in counts

    def test_detect_empty_text(self):
        """Test detecting in empty text."""
        counts = detect_special_whitespace("")

        assert counts == {}

    def test_detect_no_special_chars(self):
        """Test text without special characters."""
        text = "regular text"
        counts = detect_special_whitespace(text)

        assert counts == {}


class TestNormalizeWhitespace:
    """Test whitespace normalization."""

    def test_normalize_multiple_spaces(self):
        """Test normalizing multiple regular spaces."""
        text = "text    with    spaces"
        normalized = normalize_whitespace(text, preserve_special=False)

        assert normalized == "text with spaces"

    def test_normalize_preserve_special(self):
        """Test normalizing while preserving special characters."""
        text = "text  with\u00A0nbsp  and\tTab"
        normalized = normalize_whitespace(text, preserve_special=True)

        # Regular spaces normalized
        assert "  " not in normalized

        # Special characters preserved
        assert '\u00A0' in normalized
        assert '\t' in normalized

    def test_normalize_without_preserve(self):
        """Test normalizing without preserving special characters."""
        text = "text  with\u00A0nbsp  and\tTab"
        normalized = normalize_whitespace(text, preserve_special=False)

        # All whitespace normalized
        assert "  " not in normalized

    def test_normalize_empty_text(self):
        """Test normalizing empty text."""
        assert normalize_whitespace("") == ""
        assert normalize_whitespace(None) is None


class TestPreserveSpecialFormatting:
    """Test preservation of all special formatting."""

    def test_preserve_tabs_and_line_breaks(self):
        """Test preserving both tabs and line breaks."""
        doc = Document()
        para = doc.add_paragraph()

        original_run = para.add_run("original")
        add_tab(original_run)
        add_tab(original_run)
        add_line_break(original_run)

        translation_run = para.add_run("translation")

        preserve_special_formatting_in_run(original_run, translation_run)

        assert get_tab_count(translation_run) == 2
        assert get_line_break_count(translation_run) == 1

    def test_preserve_with_no_special_formatting(self):
        """Test preserving when original has no special formatting."""
        doc = Document()
        para = doc.add_paragraph()

        original_run = para.add_run("original")
        translation_run = para.add_run("translation")

        # Should not raise error
        preserve_special_formatting_in_run(original_run, translation_run)

        assert get_tab_count(translation_run) == 0
        assert get_line_break_count(translation_run) == 0


class TestSpecialCharsMapping:
    """Test special character mappings."""

    def test_mapping_completeness(self):
        """Test that forward and reverse mappings are complete."""
        # Every key in SPECIAL_CHARS should have reverse mapping
        for char, placeholder in SPECIAL_CHARS.items():
            assert placeholder in SPECIAL_CHARS_REVERSE
            assert SPECIAL_CHARS_REVERSE[placeholder] == char

    def test_mapping_uniqueness(self):
        """Test that placeholders are unique."""
        placeholders = list(SPECIAL_CHARS.values())
        assert len(placeholders) == len(set(placeholders))

    def test_non_breaking_space_mapping(self):
        """Test non-breaking space is mapped."""
        assert '\u00A0' in SPECIAL_CHARS
        assert SPECIAL_CHARS['\u00A0'] == '<NBSP>'

    def test_tab_mapping(self):
        """Test tab is mapped."""
        assert '\t' in SPECIAL_CHARS
        assert SPECIAL_CHARS['\t'] == '<TAB>'


class TestEdgeCases:
    """Test edge cases."""

    def test_protect_text_with_only_special_chars(self):
        """Test protecting text with only special characters."""
        text = "\t\u00A0\t\u00A0"
        protected = protect_special_characters(text)

        assert '\t' not in protected
        assert '\u00A0' not in protected
        assert '<TAB>' in protected
        assert '<NBSP>' in protected

    def test_restore_text_with_only_placeholders(self):
        """Test restoring text with only placeholders."""
        text = "<TAB><NBSP><TAB><NBSP>"
        restored = restore_special_characters(text)

        assert '<TAB>' not in restored
        assert '<NBSP>' not in restored
        assert '\t' in restored
        assert '\u00A0' in restored

    def test_protect_text_with_placeholder_like_string(self):
        """Test protecting text that contains placeholder-like strings."""
        text = "This text mentions <TAB> literally"
        protected = protect_special_characters(text)

        # Should not be affected (no actual tab character)
        assert protected == text

    def test_add_multiple_tabs_in_sequence(self):
        """Test adding multiple tabs in sequence."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("text")

        for i in range(5):
            add_tab(run)

        assert get_tab_count(run) == 5

    def test_preserve_with_empty_run(self):
        """Test preserving from empty run."""
        doc = Document()
        para = doc.add_paragraph()

        original_run = para.add_run("")
        translation_run = para.add_run("translation")

        # Should not raise error
        preserve_special_formatting_in_run(original_run, translation_run)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
