"""Tests for formatting preservation."""

import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from transit.utils.formatting import clone_run_formatting, clone_paragraph_formatting


def test_clone_run_bold():
    """Test cloning bold property."""
    doc = Document()
    para = doc.add_paragraph()

    source = para.add_run("source")
    source.bold = True

    target = para.add_run("target")
    clone_run_formatting(source, target)

    assert target.bold == True


def test_clone_run_italic():
    """Test cloning italic property."""
    doc = Document()
    para = doc.add_paragraph()

    source = para.add_run("source")
    source.italic = True

    target = para.add_run("target")
    clone_run_formatting(source, target)

    assert target.italic == True


def test_clone_run_font_size():
    """Test cloning font size."""
    doc = Document()
    para = doc.add_paragraph()

    source = para.add_run("source")
    source.font.size = Pt(16)

    target = para.add_run("target")
    clone_run_formatting(source, target)

    assert target.font.size == Pt(16)


def test_clone_paragraph_alignment():
    """Test cloning paragraph alignment."""
    doc = Document()

    source_para = doc.add_paragraph("source")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    target_para = doc.add_paragraph("target")
    clone_paragraph_formatting(source_para, target_para)

    assert target_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_clone_multiple_properties():
    """Test cloning multiple properties at once."""
    doc = Document()
    para = doc.add_paragraph()

    source = para.add_run("source")
    source.bold = True
    source.italic = True
    source.underline = True
    source.font.size = Pt(14)

    target = para.add_run("target")
    clone_run_formatting(source, target)

    assert target.bold == True
    assert target.italic == True
    assert target.underline == True
    assert target.font.size == Pt(14)
