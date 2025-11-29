"""Unit tests for hyperlink formatting preservation."""

import pytest
from docx import Document
from transit.utils.hyperlink_formatting import (
    has_hyperlink,
    get_hyperlink_url,
    get_paragraph_hyperlinks,
    add_hyperlink,
    preserve_hyperlinks_in_translation,
    _is_url_like
)


@pytest.fixture
def doc_with_hyperlinks():
    """Create document with hyperlinks."""
    doc = Document()

    # Paragraph with hyperlink
    para = doc.add_paragraph("Visit our website at ")
    add_hyperlink(para, "example.com", "https://example.com")
    para.add_run(" for more info.")

    # Paragraph with multiple hyperlinks
    para2 = doc.add_paragraph("Check ")
    add_hyperlink(para2, "Google", "https://google.com")
    para2.add_run(" and ")
    add_hyperlink(para2, "Bing", "https://bing.com")
    para2.add_run(" search engines.")

    # Regular paragraph (no hyperlinks)
    doc.add_paragraph("This is regular text without links.")

    return doc


class TestHasHyperlink:
    """Test detection of hyperlinks in runs."""

    def test_url_like_detection(self):
        """Test URL-like text detection."""
        assert _is_url_like("https://example.com") is True
        assert _is_url_like("http://test.org") is True
        assert _is_url_like("www.example.com") is True
        assert _is_url_like("test.com") is True
        assert _is_url_like("example.nl") is True
        assert _is_url_like("regular text") is False
        assert _is_url_like("no link here") is False


class TestGetHyperlinkUrl:
    """Test extraction of hyperlink URLs."""

    def test_extract_url_from_hyperlink(self, doc_with_hyperlinks):
        """Test extracting URL from hyperlink run."""
        # Get paragraph with hyperlink
        para = doc_with_hyperlinks.paragraphs[0]
        hyperlinks = get_paragraph_hyperlinks(para)

        # Should find one hyperlink
        assert len(hyperlinks) >= 1
        assert hyperlinks[0]['url'] == "https://example.com"

    def test_extract_multiple_hyperlinks(self, doc_with_hyperlinks):
        """Test extracting multiple hyperlinks from paragraph."""
        para = doc_with_hyperlinks.paragraphs[1]
        hyperlinks = get_paragraph_hyperlinks(para)

        # Should find two hyperlinks
        assert len(hyperlinks) == 2
        assert hyperlinks[0]['url'] == "https://google.com"
        assert hyperlinks[1]['url'] == "https://bing.com"


class TestGetParagraphHyperlinks:
    """Test getting all hyperlinks from paragraph."""

    def test_paragraph_with_single_hyperlink(self, doc_with_hyperlinks):
        """Test getting hyperlinks from paragraph with one link."""
        para = doc_with_hyperlinks.paragraphs[0]
        hyperlinks = get_paragraph_hyperlinks(para)

        assert len(hyperlinks) >= 1
        assert 'text' in hyperlinks[0]
        assert 'url' in hyperlinks[0]
        assert hyperlinks[0]['text'] == "example.com"

    def test_paragraph_with_multiple_hyperlinks(self, doc_with_hyperlinks):
        """Test getting hyperlinks from paragraph with multiple links."""
        para = doc_with_hyperlinks.paragraphs[1]
        hyperlinks = get_paragraph_hyperlinks(para)

        assert len(hyperlinks) == 2

        # Check first hyperlink
        assert hyperlinks[0]['text'] == "Google"
        assert hyperlinks[0]['url'] == "https://google.com"

        # Check second hyperlink
        assert hyperlinks[1]['text'] == "Bing"
        assert hyperlinks[1]['url'] == "https://bing.com"

    def test_paragraph_without_hyperlinks(self, doc_with_hyperlinks):
        """Test getting hyperlinks from paragraph without links."""
        para = doc_with_hyperlinks.paragraphs[2]
        hyperlinks = get_paragraph_hyperlinks(para)

        assert len(hyperlinks) == 0


class TestAddHyperlink:
    """Test adding hyperlinks to paragraphs."""

    def test_add_simple_hyperlink(self):
        """Test adding a simple hyperlink."""
        doc = Document()
        para = doc.add_paragraph("Click ")

        hyperlink = add_hyperlink(para, "here", "https://test.com")

        assert hyperlink is not None

        # Verify hyperlink was added
        hyperlinks = get_paragraph_hyperlinks(para)
        assert len(hyperlinks) == 1
        assert hyperlinks[0]['text'] == "here"
        assert hyperlinks[0]['url'] == "https://test.com"

    def test_add_multiple_hyperlinks(self):
        """Test adding multiple hyperlinks to same paragraph."""
        doc = Document()
        para = doc.add_paragraph()

        add_hyperlink(para, "First", "https://first.com")
        para.add_run(" and ")
        add_hyperlink(para, "Second", "https://second.com")

        hyperlinks = get_paragraph_hyperlinks(para)
        assert len(hyperlinks) == 2

    def test_add_url_as_text_and_link(self):
        """Test adding URL both as display text and link."""
        doc = Document()
        para = doc.add_paragraph()

        url = "https://example.org"
        add_hyperlink(para, url, url)

        hyperlinks = get_paragraph_hyperlinks(para)
        assert len(hyperlinks) == 1
        assert hyperlinks[0]['text'] == url
        assert hyperlinks[0]['url'] == url


class TestPreserveHyperlinks:
    """Test preservation of hyperlinks in translation."""

    def test_preserve_logs_hyperlinks(self, doc_with_hyperlinks, caplog):
        """Test that preserve function logs hyperlinks."""
        import logging
        caplog.set_level(logging.INFO)

        original_para = doc_with_hyperlinks.paragraphs[0]
        translation_para = doc_with_hyperlinks.add_paragraph("Bezoek onze website")

        preserve_hyperlinks_in_translation(original_para, translation_para)

        # Should log information about hyperlinks
        assert "hyperlink" in caplog.text.lower()

    def test_preserve_with_no_hyperlinks(self, caplog):
        """Test preservation when paragraph has no hyperlinks."""
        import logging
        caplog.set_level(logging.INFO)

        doc = Document()
        original_para = doc.add_paragraph("No links here")
        translation_para = doc.add_paragraph("Geen links hier")

        # Should not raise error
        preserve_hyperlinks_in_translation(original_para, translation_para)

        # Should not log anything significant
        # (may have debug logs, but not info about hyperlinks)

    def test_preserve_with_multiple_hyperlinks(self, doc_with_hyperlinks, caplog):
        """Test preservation with multiple hyperlinks."""
        import logging
        caplog.set_level(logging.INFO)

        original_para = doc_with_hyperlinks.paragraphs[1]  # Has 2 hyperlinks
        translation_para = doc_with_hyperlinks.add_paragraph("Controleer zoekmachines")

        preserve_hyperlinks_in_translation(original_para, translation_para)

        # Should log information about both hyperlinks
        log_text = caplog.text.lower()
        assert "hyperlink" in log_text or "hyperlinks" in log_text


class TestEdgeCases:
    """Test edge cases in hyperlink handling."""

    def test_empty_paragraph_no_hyperlinks(self):
        """Test empty paragraph has no hyperlinks."""
        doc = Document()
        para = doc.add_paragraph("")

        hyperlinks = get_paragraph_hyperlinks(para)
        assert len(hyperlinks) == 0

    def test_whitespace_only_paragraph(self):
        """Test whitespace-only paragraph has no hyperlinks."""
        doc = Document()
        para = doc.add_paragraph("   ")

        hyperlinks = get_paragraph_hyperlinks(para)
        assert len(hyperlinks) == 0

    def test_add_hyperlink_with_special_characters(self):
        """Test adding hyperlink with special characters in text."""
        doc = Document()
        para = doc.add_paragraph()

        special_text = "Click here! (special)"
        add_hyperlink(para, special_text, "https://test.com")

        hyperlinks = get_paragraph_hyperlinks(para)
        assert len(hyperlinks) == 1
        assert hyperlinks[0]['text'] == special_text

    def test_add_hyperlink_with_long_url(self):
        """Test adding hyperlink with very long URL."""
        doc = Document()
        para = doc.add_paragraph()

        long_url = "https://example.com/very/long/path/with/many/segments?param1=value1&param2=value2"
        add_hyperlink(para, "Click", long_url)

        hyperlinks = get_paragraph_hyperlinks(para)
        assert len(hyperlinks) == 1
        assert hyperlinks[0]['url'] == long_url

    def test_preserve_with_empty_original(self):
        """Test preserve with empty original paragraph."""
        doc = Document()
        original = doc.add_paragraph("")
        translation = doc.add_paragraph("Translation")

        # Should not raise error
        preserve_hyperlinks_in_translation(original, translation)

    def test_preserve_with_empty_translation(self):
        """Test preserve with empty translation paragraph."""
        doc = Document()
        original = doc.add_paragraph("Text with ")
        add_hyperlink(original, "link", "https://test.com")
        translation = doc.add_paragraph("")

        # Should not raise error
        preserve_hyperlinks_in_translation(original, translation)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
