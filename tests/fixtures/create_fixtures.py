"""Script to create test fixture DOCX files for testing."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_simple_document():
    """Create simple document with plain text."""
    doc = Document()
    doc.add_paragraph("Dit is een simpele paragraaf.")
    doc.add_paragraph("Dit is een tweede paragraaf.")
    doc.add_paragraph("Dit is een derde paragraaf.")
    return doc


def create_formatted_document():
    """Create document with complex formatting."""
    doc = Document()

    # Paragraph with bold
    para1 = doc.add_paragraph()
    run1 = para1.add_run("Dit is vetgedrukt.")
    run1.bold = True

    # Paragraph with italic
    para2 = doc.add_paragraph()
    run2 = para2.add_run("Dit is schuin.")
    run2.italic = True

    # Paragraph with underline
    para3 = doc.add_paragraph()
    run3 = para3.add_run("Dit is onderstreept.")
    run3.underline = True

    # Paragraph with mixed formatting
    para4 = doc.add_paragraph()
    run4a = para4.add_run("Normale tekst ")
    run4b = para4.add_run("vetgedrukt")
    run4b.bold = True
    run4c = para4.add_run(" en ")
    run4d = para4.add_run("schuin")
    run4d.italic = True
    run4e = para4.add_run(" tekst.")

    # Paragraph with custom font size
    para5 = doc.add_paragraph()
    run5 = para5.add_run("Grote tekst")
    run5.font.size = Pt(24)

    # Paragraph with color
    para6 = doc.add_paragraph()
    run6 = para6.add_run("Gekleurde tekst")
    run6.font.color.rgb = RGBColor(255, 0, 0)  # Red

    # Centered paragraph
    para7 = doc.add_paragraph("Gecentreerde tekst")
    para7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Right-aligned paragraph
    para8 = doc.add_paragraph("Rechts uitgelijnd")
    para8.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def create_table_document():
    """Create document with tables."""
    doc = Document()

    doc.add_paragraph("Document met tabellen")

    # Simple 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # Fill table with data
    for i in range(3):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = f"Cel ({i}, {j})"

    doc.add_paragraph("Tabel hierboven")

    return doc


def create_merged_cells_document():
    """Create document with merged cells in table."""
    doc = Document()

    doc.add_paragraph("Tabel met samengevoegde cellen")

    # Create table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # Merge cells horizontally (row 0, cols 0-2)
    cell_a = table.cell(0, 0)
    cell_b = table.cell(0, 2)
    cell_a.merge(cell_b)
    cell_a.text = "Samengevoegde cel (horizontaal)"

    # Merge cells vertically (rows 1-2, col 0)
    cell_c = table.cell(1, 0)
    cell_d = table.cell(2, 0)
    cell_c.merge(cell_d)
    cell_c.text = "Samengevoegd (verticaal)"

    # Fill remaining cells
    table.cell(1, 1).text = "Cel (1,1)"
    table.cell(1, 2).text = "Cel (1,2)"
    table.cell(2, 1).text = "Cel (2,1)"
    table.cell(2, 2).text = "Cel (2,2)"

    return doc


def create_header_footer_document():
    """Create document with headers and footers."""
    doc = Document()

    # Add header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Dit is de header"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add body content
    doc.add_paragraph("Body paragraaf 1")
    doc.add_paragraph("Body paragraaf 2")
    doc.add_paragraph("Body paragraaf 3")

    # Add footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Dit is de footer - Pagina"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def create_complex_document():
    """Create complex document with multiple elements."""
    doc = Document()

    # Header
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Complexe Document Header"

    # Title
    title = doc.add_paragraph("Hoofdtitel van Document")
    title.style = 'Heading 1'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_paragraph("Dit is een complexe document met verschillende elementen.")

    # Section 1
    doc.add_paragraph("Sectie 1: Tekst met Formatting", style='Heading 2')

    para1 = doc.add_paragraph()
    run1a = para1.add_run("Dit is ")
    run1b = para1.add_run("belangrijke")
    run1b.bold = True
    run1c = para1.add_run(" informatie over ")
    run1d = para1.add_run("specifieke")
    run1d.italic = True
    run1e = para1.add_run(" onderwerpen.")

    # Section 2 with table
    doc.add_paragraph("Sectie 2: Data Tabel", style='Heading 2')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Aantal'
    hdr_cells[2].text = 'Prijs'

    # Data rows
    data = [
        ('Product A', '10', '€50'),
        ('Product B', '25', '€75'),
        ('Product C', '15', '€60')
    ]

    for i, (item, aantal, prijs) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = item
        row_cells[1].text = aantal
        row_cells[2].text = prijs

    # Section 3
    doc.add_paragraph("Sectie 3: Conclusie", style='Heading 2')
    doc.add_paragraph("Dit document demonstreert verschillende DOCX features.")

    # Footer
    footer = section.footer
    footer.paragraphs[0].text = "Pagina Footer - Confidentieel"

    return doc


def create_special_characters_document():
    """Create document with special characters."""
    doc = Document()

    doc.add_paragraph("Document met speciale karakters")

    # Non-breaking space
    doc.add_paragraph("Tekst met\u00A0non-breaking\u00A0space")

    # Tab character
    doc.add_paragraph("Tekst met\ttab\tkarakters")

    # Multiple spaces
    doc.add_paragraph("Tekst    met    meerdere    spaties")

    # Line break within paragraph
    para = doc.add_paragraph("Tekst met")
    para.add_run("\n")
    para.add_run("line break")

    # Various punctuation
    doc.add_paragraph("Tekst met: aanhalingstekens 'test', dubbele \"test\", en uitroepteken!")
    doc.add_paragraph("Vraag? En antwoord.")
    doc.add_paragraph("Ellipsis... en gedachtestreepje - test.")

    return doc


def create_empty_elements_document():
    """Create document with empty and whitespace elements."""
    doc = Document()

    doc.add_paragraph("Document met lege elementen")

    # Empty paragraph
    doc.add_paragraph("")

    # Whitespace-only paragraph
    doc.add_paragraph("   ")

    # Paragraph with text
    doc.add_paragraph("Normale tekst")

    # Another empty
    doc.add_paragraph("")

    # Table with empty cells
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tekst"
    # Other cells stay empty

    return doc


def create_abbreviations_document():
    """Create document with common Dutch abbreviations."""
    doc = Document()

    doc.add_paragraph("Document met afkortingen")

    doc.add_paragraph("Dit is een test m.b.t. de vertaling van afkortingen.")
    doc.add_paragraph("Bijvoorbeeld (b.v.) moeten deze correct vertaald worden.")
    doc.add_paragraph("Dr. Smith en Mevr. Johnson waren aanwezig.")
    doc.add_paragraph("De vergadering was o.a. over budgetten.")
    doc.add_paragraph("Het bedrijf is gevestigd in de VS (Verenigde Staten).")
    doc.add_paragraph("De prijs is ca. €100 excl. BTW.")

    return doc


def create_list_document():
    """Create document with bulleted and numbered lists."""
    doc = Document()

    doc.add_paragraph("Document met lijsten")

    # Bulleted list
    doc.add_paragraph("Ongenummerde lijst:")
    doc.add_paragraph("Eerste item", style='List Bullet')
    doc.add_paragraph("Tweede item", style='List Bullet')
    doc.add_paragraph("Derde item", style='List Bullet')

    doc.add_paragraph("")  # Spacer

    # Numbered list
    doc.add_paragraph("Genummerde lijst:")
    doc.add_paragraph("Eerste stap", style='List Number')
    doc.add_paragraph("Tweede stap", style='List Number')
    doc.add_paragraph("Derde stap", style='List Number')

    doc.add_paragraph("")  # Spacer

    # Nested list (approximation - manual indentation)
    doc.add_paragraph("Geneste lijst:")
    doc.add_paragraph("Hoofditem 1", style='List Bullet')
    doc.add_paragraph("Sub-item 1.1", style='List Bullet 2')
    doc.add_paragraph("Sub-item 1.2", style='List Bullet 2')
    doc.add_paragraph("Hoofditem 2", style='List Bullet')

    return doc


def create_hyperlink_document():
    """Create document with hyperlinks."""
    from transit.utils.hyperlink_formatting import add_hyperlink

    doc = Document()

    doc.add_paragraph("Document met hyperlinks")

    # Paragraph with single hyperlink
    para1 = doc.add_paragraph("Bezoek onze website op ")
    add_hyperlink(para1, "https://example.com", "https://example.com")
    para1.add_run(" voor meer informatie.")

    # Paragraph with text hyperlink
    para2 = doc.add_paragraph("Klik ")
    add_hyperlink(para2, "hier", "https://google.com")
    para2.add_run(" om te zoeken.")

    # Paragraph with multiple hyperlinks
    para3 = doc.add_paragraph("Bekijk ")
    add_hyperlink(para3, "Google", "https://google.com")
    para3.add_run(" en ")
    add_hyperlink(para3, "Bing", "https://bing.com")
    para3.add_run(" voor zoeken.")

    # Email hyperlink
    para4 = doc.add_paragraph("Contact: ")
    add_hyperlink(para4, "info@example.com", "mailto:info@example.com")

    return doc


def create_nested_tables_document():
    """Create document with nested tables."""
    doc = Document()

    doc.add_paragraph("Document met geneste tabellen")

    # Main table
    main_table = doc.add_table(rows=3, cols=2)
    main_table.style = 'Table Grid'

    # Fill some cells with regular text
    main_table.cell(0, 0).text = "Rij 1, Kolom 1"
    main_table.cell(0, 1).text = "Rij 1, Kolom 2"
    main_table.cell(1, 0).text = "Rij 2, Kolom 1"

    # Add nested table in cell (1, 1)
    cell_with_nested = main_table.cell(1, 1)
    cell_with_nested.text = "Deze cel bevat een geneste tabel:"

    # Create nested table within the cell
    nested_table = cell_with_nested.add_table(rows=2, cols=2)
    nested_table.style = 'Table Grid'

    nested_table.cell(0, 0).text = "Genest A1"
    nested_table.cell(0, 1).text = "Genest A2"
    nested_table.cell(1, 0).text = "Genest B1"
    nested_table.cell(1, 1).text = "Genest B2"

    # Continue with main table
    main_table.cell(2, 0).text = "Rij 3, Kolom 1"
    main_table.cell(2, 1).text = "Rij 3, Kolom 2"

    doc.add_paragraph("Tabel met geneste structuur hierboven")

    return doc


def main():
    """Create all test fixtures."""
    fixtures_dir = os.path.dirname(os.path.abspath(__file__))

    fixtures = {
        'simple.docx': create_simple_document,
        'formatted.docx': create_formatted_document,
        'table.docx': create_table_document,
        'merged_cells.docx': create_merged_cells_document,
        'header_footer.docx': create_header_footer_document,
        'complex.docx': create_complex_document,
        'special_chars.docx': create_special_characters_document,
        'empty_elements.docx': create_empty_elements_document,
        'abbreviations.docx': create_abbreviations_document,
        'lists.docx': create_list_document,
        'hyperlinks.docx': create_hyperlink_document,
        'nested_tables.docx': create_nested_tables_document,
    }

    print("Creating test fixtures...")
    for filename, create_func in fixtures.items():
        filepath = os.path.join(fixtures_dir, filename)
        doc = create_func()
        doc.save(filepath)
        print(f"  Created: {filename}")

    print(f"\nAll fixtures created in: {fixtures_dir}")
    print("\nFixtures:")
    for filename in fixtures.keys():
        print(f"  - {filename}")


if __name__ == "__main__":
    main()
